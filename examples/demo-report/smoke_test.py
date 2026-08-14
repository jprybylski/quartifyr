#!/usr/bin/env python3
"""End-to-end smoke test for the quartifyr demo report.

Runs the real two-pass pipeline (Rscript render.R --final, which drives
Quarto + reportifyr exactly as a report author would) and asserts on the
resulting docx: no leftover reportifyr magic strings, the expected tables/
images/abbreviations are present, and appendix numbering resolved.

Requires the full toolchain: R (with this repo's renv-managed packages),
Quarto, and the Python venv. Run from anywhere:

    python3 examples/demo-report/smoke_test.py
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import docx
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_DIR = Path(__file__).resolve().parent
REPO_ROOT = PROJECT_DIR.parent.parent


def _iter_paragraph_text(document: docx.Document):
    for p in document.element.body.iter(qn("w:p")):
        texts = p.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts)
        if text.strip():
            yield text


def main() -> int:
    if shutil.which("Rscript") is None:
        print("SKIP: Rscript not found on PATH", file=sys.stderr)
        return 0
    if shutil.which("quarto") is None:
        print("SKIP: quarto not found on PATH", file=sys.stderr)
        return 0

    # Quarto's extension loader doesn't follow symlinks, so this demo
    # carries a physical copy of _extensions/quartifyr/ rather than a
    # symlink to the canonical one at the repo root -- which means it can
    # silently drift out of sync with real fixes (this happened for real:
    # a Heading1-vs-"Heading 1" style-ID bugfix landed in the canonical
    # copy but the demo kept rendering with the stale, buggy version until
    # caught). Fail fast, before wasting time on a render that would only
    # prove the *old* code works.
    sync_check = subprocess.run(
        ["python3", str(REPO_ROOT / "scripts" / "sync_demo_extension.py"), "--check"],
        capture_output=True,
        text=True,
    )
    if sync_check.returncode != 0:
        print("FAIL: demo's _extensions/quartifyr/ copy has drifted from the canonical one", file=sys.stderr)
        print(sync_check.stdout, file=sys.stderr)
        print(sync_check.stderr, file=sys.stderr)
        return 1

    # inst/templates/org-reference.docx is committed (see scripts/
    # check_template_freshness.py's docstring for why) -- fail fast if it's
    # drifted from inst/python/styles/default.yaml, before wasting time on a
    # render that would only prove the *old* styling still works.
    template_check = subprocess.run(
        ["python3", str(REPO_ROOT / "scripts" / "check_template_freshness.py"), "--check"],
        capture_output=True,
        text=True,
    )
    if template_check.returncode != 0:
        print("FAIL: inst/templates/org-reference.docx has drifted from inst/python/styles/default.yaml", file=sys.stderr)
        print(template_check.stdout, file=sys.stderr)
        print(template_check.stderr, file=sys.stderr)
        return 1

    print("Running Rscript render.R --final ...")
    result = subprocess.run(
        ["Rscript", "render.R", "--final"],
        cwd=PROJECT_DIR,
        capture_output=True,
        text=True,
        timeout=180,
    )
    if result.returncode != 0:
        print("FAIL: render.R exited non-zero", file=sys.stderr)
        print(result.stdout, file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return 1

    final_docx = PROJECT_DIR / "report" / "final" / "report-final.docx"
    if not final_docx.exists():
        print(f"FAIL: expected output not found: {final_docx}", file=sys.stderr)
        return 1

    document = docx.Document(str(final_docx))
    all_text = list(_iter_paragraph_text(document))
    joined = " | ".join(all_text)

    checks: list[tuple[str, bool]] = []

    # quartifyr's own pass-1 output (Quarto render + apply-layout, before
    # reportifyr's pass-2 fill touches anything) must never leave an
    # unpaired w:bookmarkStart/w:bookmarkEnd -- confirmed, via a real bug
    # reported against this exact demo, that reportifyr's own
    # `remove_bookmarks()` step (python, docx_utils.py) matches
    # bookmarkEnd elements to remove *purely by numeric w:id*, without also
    # checking that the matching bookmarkStart's name carries reportifyr's
    # own "fp_" footnote-bookmark prefix. reportifyr assigns its own
    # footnote-bookmark ids from the footnoted paragraph's index in
    # doc.paragraphs -- a small integer with no relationship to pandoc's
    # own citeproc/heading bookmark ids -- so on some documents (this one
    # included) it coincidentally collides with an unrelated bookmark
    # (here, a citeproc `ref-<key>` bookmark), and remove_bookmarks() then
    # deletes *both* same-numbered bookmarkEnd elements, silently
    # corrupting whichever unrelated bookmark lost its close tag -- e.g. a
    # citation's `link-citations: true` hyperlink pointing at a bookmark
    # that's missing its end. This is a reportifyr bug, not quartifyr's;
    # this check only proves quartifyr's own shell is never the source.
    shell_docx = PROJECT_DIR / "report" / "shell" / "report.docx"
    with zipfile.ZipFile(shell_docx) as z:
        shell_xml = z.read("word/document.xml").decode("utf-8")
    shell_start_ids = set(re.findall(r'<w:bookmarkStart w:id="(\d+)"', shell_xml))
    shell_end_ids = set(re.findall(r'<w:bookmarkEnd w:id="(\d+)"', shell_xml))
    checks.append(
        (
            "quartifyr's own shell (pre-reportifyr) has no unpaired bookmarks",
            shell_start_ids == shell_end_ids,
        )
    )

    checks.append(("no leftover {rpfy}: magic strings", "{rpfy}:" not in joined))
    checks.append(("at least 7 tables present (title info + signatures + PK summary + demographics + abbreviations; synopsis is plain paragraphs, not a table)", len(document.tables) >= 7))

    # build_template.py's reference-doc sets <w:updateFields w:val="true"/>
    # in settings.xml so Word auto-recalculates every field (TOC, SEQ, REF,
    # PAGE, ...) the moment a human opens the delivered file -- no manual
    # "select all, F9", no dependency on recalculate-fields' headless
    # LibreOffice automation. Confirmed here to survive the full pipeline
    # (Quarto's reference-doc handling, apply-layout, reportifyr) into the
    # actual delivered docx, not just the reference-doc itself.
    update_fields = document.settings.element.find(qn("w:updateFields"))
    checks.append(
        (
            "delivered docx auto-recalculates fields on open in Word (w:updateFields survived the full pipeline)",
            update_fields is not None and update_fields.get(qn("w:val")) == "true",
        )
    )

    with zipfile.ZipFile(final_docx) as z:
        images = [n for n in z.namelist() if n.startswith("word/media/")]
    checks.append(("figure embedded as a real image", len(images) >= 1))

    checks.append(("title page title rendered", "Population Pharmacokinetics of Theophylline" in joined))
    checks.append(("document-status stamp rendered", "DRAFT" in joined or "FINAL" in joined))
    checks.append(("\\gls{PK} resolved to a real definition", "pharmacokinetic (pk)" in joined.lower()))

    # title_page.lua's "Title Page" ToC-entry marker -- an actual
    # Heading1-styled paragraph (real outline level, picked up by Word's
    # native ToC field automatically, no apply-layout step needed), made
    # invisible on the page itself via ordinary formatting (1pt size,
    # white color) rather than any "hidden" flag -- confirmed via direct
    # testing in real Word that both `<w:vanish/>` paragraphs and `TC`
    # fields (two earlier attempts) reliably showed *something* on the
    # title page regardless. See title_page.lua's file-header comment.
    checks.append(("'Title Page' ToC-entry marker present", "Title Page" in joined))
    with zipfile.ZipFile(final_docx) as z:
        document_xml_for_toc = z.read("word/document.xml").decode("utf-8")
    checks.append(
        (
            "'Title Page' marker is actually Heading1-styled, tiny, and white",
            re.search(
                r'<w:pStyle w:val="Heading1"/>.*?<w:sz w:val="2"/><w:szCs w:val="2"/><w:color w:val="FFFFFF"/></w:rPr><w:t[^>]*>Title Page',
                document_xml_for_toc,
                re.DOTALL,
            )
            is not None,
        )
    )
    # <w:vanish/> would exclude the paragraph from Word's ToC outline scan
    # entirely (confirmed directly in real Word) -- a regression here would
    # silently make the "Title Page" entry disappear from the ToC again.
    title_page_marker_match = re.search(
        r'<w:pStyle w:val="Heading1"/>.*?<w:t[^>]*>Title Page', document_xml_for_toc, re.DOTALL
    )
    checks.append(
        (
            "'Title Page' marker does not use <w:vanish/> (would break ToC inclusion)",
            title_page_marker_match is not None and "<w:vanish/>" not in title_page_marker_match.group(0),
        )
    )

    # Contributors and Approvers now share one "Signatures" page/ToC entry
    # (see signature_page.lua) rather than two separate Heading-1 sections.
    checks.append(("signature page uses a single 'Signatures' heading", "Signatures" in joined))
    checks.append(("Contributors sub-label rendered within Signatures", "Contributors" in joined))
    checks.append(("Approvers sub-label rendered within Signatures", "Approvers" in joined))

    def _table_row_texts(table) -> list[list[str]]:
        return [[cell.text for cell in row.cells] for row in table.rows]

    all_rows = [row for table in document.tables for row in _table_row_texts(table)]
    checks.append(
        ("abbreviations table has a PK row", any("PK" in row and "Pharmacokinetic" in row for row in all_rows))
    )
    checks.append(("appendix auto-lettered", "Appendix A: Statistical Methods" in joined))

    # Second appendix: Quarto's own native `{{< include >}}` shortcode
    # nested in a fenced code block (report.qmd), not a quartifyr
    # shortcode -- embeds scripts/01_analysis.R's real content with real
    # syntax highlighting. Both appendix headings show the same cached
    # "A" (see appendix.lua's file-header comment: the SEQ Appendix field's
    # displayed letter is a static placeholder pending a Word field
    # recalculation, same as every other {{< appendix >}} heading in this
    # doc -- not a regression here).
    checks.append(("second appendix (embedded analysis script) present", "Appendix A: Analysis Code" in joined))
    checks.append(
        (
            "embedded script's real content present, not an empty/failed include",
            "aggregate" in joined and "flextable(theoph_demographics)" in joined,
        )
    )
    with zipfile.ZipFile(final_docx) as z:
        final_xml = z.read("word/document.xml").decode("utf-8")
    checks.append(
        (
            "embedded script is syntax-highlighted (pandoc token character styles present)",
            "FunctionTok" in final_xml and "CommentTok" in final_xml,
        )
    )

    # bibliography: a `bibliography: references.bib` shell .qmd is plain
    # Quarto/pandoc citeproc support -- no quartifyr-specific code involved
    # in generating the citations/reference list themselves. What this
    # extension DOES add: a "Bibliography" paragraph style and a
    # "Hyperlink" character style on the reference-doc (pandoc's docx
    # writer applies these by pStyle/rStyle, but neither is one of
    # python-docx's bundled built-in styles, so build_template.py has to
    # define them -- see _style_bibliography()/_style_hyperlink()), and
    # bibliography.lua's `csl:`/`link-citations:` defaults (this
    # extension's bundled NLM/Vancouver citation-sequence-brackets style,
    # with in-text citations hyperlinked to their bibliography entry, used
    # whenever a project doesn't set its own -- see that filter's
    # file-header comment).
    checks.append(
        (
            "in-text citation resolved in NLM/Vancouver bracketed numeric style (default csl, not pandoc's author-date default)",
            "built-in Theoph dataset [1]," in joined and "(Boeckmann et al. 1994)" not in joined,
        )
    )
    # Plain "References", not numbered like Introduction/Results ("1. ",
    # "2. ", literal tabs baked in by number-sections) -- it uses the same
    # `custom-style="Heading 1"` Div idiom as Synopsis/Abbreviations above
    # it in this file, deliberately NOT an actual pandoc Header, so
    # number-sections never numbers it (see report.qmd's comment for why a
    # real `# References {.unnumbered}` heading doesn't work cleanly:
    # quarto-plus's header.lua still inserts its tab, leaving one with no
    # number before it).
    checks.append(("References heading rendered, not numbered", any(t.strip() == "References" for t in all_text)))
    checks.append(
        (
            "bibliography entries rendered in NLM author-initials format",
            any(t.startswith("1. \tBoeckmann AJ, Sheiner LB, Beal SL.") for t in all_text)
            and any(t.startswith("2. \tPinheiro JC, Bates DM.") for t in all_text),
        )
    )

    # The shell .qmd places an explicit `::: {#refs}` Div right after the
    # body and before the appendices -- citeproc fills that Div in place
    # rather than appending the bibliography to the very end of the
    # document (its default behavior absent an explicit #refs Div), which
    # would otherwise land it after all appendix content instead.
    references_idx = next(i for i, t in enumerate(all_text) if t.strip() == "References")
    appendix_idx = next(i for i, t in enumerate(all_text) if "Appendix A: Statistical Methods" in t)
    checks.append(("References section populates before the appendices, not after", references_idx < appendix_idx))
    checks.append(
        ("PK summary table header present", any({"Cmax", "Cmin"}.issubset(set(row)) for row in all_rows))
    )
    checks.append(
        (
            "participant demographics table header present",
            any({"Weight (kg)", "Dose (mg/kg)"}.issubset(set(row)) for row in all_rows),
        )
    )

    # The "style bleed" contrast scripts/01_analysis.R's comment describes:
    # pk_summary started life as a plain data frame (.csv), so reportifyr's
    # add_tables() ran its own format_flextable() on it, which hardcodes
    # Arial Narrow regardless of this doc's actual body font (Times New
    # Roman, per inst/python/styles/default.yaml). The demographics table
    # started life as an already-built `flextable` object (.rds), which
    # add_tables() inserts as-is instead of reformatting it. Checked here
    # against the real rendered docx, not just asserted by construction.
    def _table_with_header_cell(needle: str):
        for t in document.tables:
            if t.rows and any(needle in cell.text for cell in t.rows[0].cells):
                return t
        return None

    def _first_body_run_font(table) -> str | None:
        if table is None or len(table.rows) < 2:
            return None
        run = table.rows[1].cells[0].paragraphs[0].runs[0]
        return run.font.name

    pk_table_font = _first_body_run_font(_table_with_header_cell("Cmax"))
    demographics_table_font = _first_body_run_font(_table_with_header_cell("Weight (kg)"))
    checks.append(
        (
            "PK summary table body renders in reportifyr's own hardcoded Arial Narrow (plain-data-frame table always gets reformatted, regardless of this doc's Times New Roman body font)",
            pk_table_font == "Arial Narrow",
        )
    )
    checks.append(
        (
            "demographics table body renders in its own hand-styled Times New Roman (pre-built flextable .rds is inserted as-is, matching this doc's body font)",
            demographics_table_font == "Times New Roman",
        )
    )

    checks.append(("synopsis section rendered", "Synopsis" in joined))

    # Front-matter section labels (Synopsis, Table of Contents, List of
    # Figures/Tables, Abbreviations) use pandoc's `custom-style="Heading
    # 1"` Div attribute to get the real Heading 1 look without becoming a
    # real numbered heading. This matches by *style name* ("Heading 1",
    # with a space) -- get that backwards (the style *ID* "Heading1" has
    # no space, the opposite convention from raw-OOXML pStyle references
    # elsewhere in this extension) and pandoc silently fabricates a new,
    # blank style with that literal name instead of erroring, so the text
    # renders as plain, unstyled body text with no indication anything's
    # wrong. Confirmed by making this exact mistake for real -- see
    # _extensions/quartifyr/README.md's pStyle gotcha section.
    checks.append(
        (
            "front-matter section labels use the real bold Heading 1 style, not a blank phantom one",
            "<w:b/>" in document.styles["Heading 1"].element.xml,
        )
    )

    # Synopsis is plain flowing paragraphs (label, then value line(s)), not
    # a table -- deliberately: reportifyr's auto-generated footnote for a
    # figure embedded in a table cell groups per *table element* and lands
    # after the whole table, not under the specific figure. Plain
    # body-level paragraphs get reportifyr's already-correct per-figure
    # footnote placement instead (the same mechanism the body's own
    # Figure 1 already uses without issue) -- see synopsis.lua's own
    # comment for the full reasoning trail.
    body_paragraphs = list(document.element.body.iter(qn("w:p")))
    body_paragraph_texts = ["".join(t.text or "" for t in p.findall(".//" + qn("w:t"))) for p in body_paragraphs]

    # report.qmd sets synopsis-style: inline -- a row with a single
    # non-image value line (Title, Objectives, Methods) renders as one
    # "Label:  value" paragraph, label and value in separate runs so the
    # label alone can carry the bold/larger "Synopsis Inline Label"
    # character style (see build_template.py).
    title_row_idx = next(
        (i for i, t in enumerate(body_paragraph_texts) if t == "Title:  Population Pharmacokinetics of Theophylline"),
        None,
    )
    checks.append(("synopsis title row has the real report title", title_row_idx is not None))

    title_row_rstyle = None
    if title_row_idx is not None:
        title_row_rstyle = body_paragraphs[title_row_idx].find(".//" + qn("w:rStyle"))
    checks.append(
        (
            "synopsis-style: inline's label run uses the bold, larger 'Synopsis Inline Label' character style",
            title_row_rstyle is not None and title_row_rstyle.get(qn("w:val")) == "SynopsisInlineLabel",
        )
    )

    # Regression check: LibreOffice doesn't reliably give a run's
    # referenced character style priority over its paragraph style's own
    # bold -- confirmed by rendering this exact row and finding the label
    # plain, the value bold (backwards from both styles' definitions).
    # Both runs need an explicit direct <w:b/>/<w:b w:val="0"/>, not just
    # style references, to render correctly regardless of that cascade.
    title_row_runs = body_paragraphs[title_row_idx].findall(qn("w:r")) if title_row_idx is not None else []
    title_label_run_b = title_row_runs[0].find(qn("w:rPr") + "/" + qn("w:b")) if len(title_row_runs) > 0 else None
    title_value_run_b = title_row_runs[1].find(qn("w:rPr") + "/" + qn("w:b")) if len(title_row_runs) > 1 else None
    checks.append(("inline's label run has an explicit direct <w:b/>, not just the rStyle reference", title_label_run_b is not None and title_label_run_b.get(qn("w:val")) != "0"))
    checks.append(("inline's value run has an explicit direct <w:b w:val=\"0\"/> overriding the paragraph style's bold", title_value_run_b is not None and title_value_run_b.get(qn("w:val")) == "0"))

    # Results' first value line is plain text (and more than one
    # sentence, to prove the merge isn't sentence-limited -- it's just
    # whatever the first YAML list item contains, verbatim), so the
    # label still merges into it even though two more lines (more text,
    # then an embedded image) follow -- only the *first* line's shape
    # decides mergeability (see synopsis.lua's add_row).
    results_merge_idx = next(
        (
            i
            for i, t in enumerate(body_paragraph_texts)
            if t
            == "Results:  Peak concentrations and inter-participant variability are summarized in Table 1 and Figure 1. "
            "Individual profiles show wide variability in both peak concentration and time to peak across "
            "participants, consistent with typical oral absorption variability for this compound."
        ),
        None,
    )
    checks.append(("synopsis-style: inline merges the label into Results' full (multi-sentence) first line", results_merge_idx is not None))

    results_rstyle = None
    if results_merge_idx is not None:
        results_rstyle = body_paragraphs[results_merge_idx].find(".//" + qn("w:rStyle"))
    checks.append(
        (
            "Results' merged label run uses the bold/larger 'Synopsis Inline Label' character style too",
            results_rstyle is not None and results_rstyle.get(qn("w:val")) == "SynopsisInlineLabel",
        )
    )

    # The row's second (continuation) line still follows underneath,
    # flush -- no repeated label, no indent (indentation is a
    # definition-list trait only; inline never indents).
    continuation_ind = None
    if results_merge_idx is not None:
        continuation_ppr = body_paragraphs[results_merge_idx + 1].find(qn("w:pPr"))
        if continuation_ppr is not None:
            continuation_ind = continuation_ppr.find(qn("w:ind"))
    checks.append(
        (
            "Results' second value line follows underneath flush (w:ind left=0 override present, not inherited indent)",
            continuation_ind is not None and continuation_ind.get(qn("w:left")) == "0",
        )
    )

    checks.append(("logo image embedded on title page", any("png" in n.lower() for n in images)))
    checks.append(("address rendered on title page", "Raleigh, NC" in joined))

    # report_number is a plain YAML anchor (&report_number), aliased
    # (*report_number) into this title-page-extra row rather than
    # written out a second time -- see report.qmd's comment. Both
    # occurrences below prove the *same* written value reached the title
    # page's title-page-extra row and the header's header-format
    # placeholder (checked separately by header_expected further down).
    checks.append(("title-page-extra 'Report Number' row rendered from the anchored report_number", "Report Number | RPT-2026-014" in joined))

    checks.append(("synopsis value supports multi-line text", "demonstrating multi-line values with an embedded figure" in joined))
    checks.append(("synopsis figure embedded as a real image", len(images) >= 2))

    # The synopsis figure's magic string is the only one with a <width:...>
    # arg (this filter always sets one) -- distinct from the body's own
    # Figure 1 (no width arg, uses the artifact's original size) and the
    # title-page logo (not a reportifyr figure at all), so its alt text is
    # how to find it specifically among all the document's drawings.
    synopsis_fig_idx = None
    synopsis_fig_alt = ""
    for i, p in enumerate(body_paragraphs):
        for d in p.findall(".//" + qn("w:drawing")):
            for dp in d.findall(".//" + qn("wp:docPr")):
                descr = dp.get("descr") or ""
                if "<width:" in descr:
                    synopsis_fig_idx = i
                    synopsis_fig_alt = descr

    checks.append(("synopsis figure embedded as a real body-level image (not inside a table)", synopsis_fig_idx is not None))
    checks.append(
        (
            "synopsis figure carries reportifyr's own provenance metadata (a content hash) as alt text",
            "hash:" in synopsis_fig_alt,
        )
    )
    checks.append(
        (
            "synopsis figure's footnote lands immediately after the figure itself, not after an unrelated table",
            synopsis_fig_idx is not None
            and synopsis_fig_idx + 1 < len(body_paragraph_texts)
            and body_paragraph_texts[synopsis_fig_idx + 1].startswith("[Source:"),
        )
    )

    with zipfile.ZipFile(final_docx) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
    seq_figure_count = document_xml.count("SEQ Figure")
    checks.append(
        (
            "synopsis figure excluded from the List of Figures (only Figure 1's own SEQ Figure field exists)",
            seq_figure_count == 1,
        )
    )

    checks.append(
        (
            "bibliography entries use the reference-doc's own 'Bibliography' paragraph style",
            '<w:pStyle w:val="Bibliography"' in document_xml,
        )
    )

    # link-citations: true (bibliography.lua's default) makes each in-text
    # citation a real internal hyperlink to its bibliography entry -- one
    # <w:hyperlink> anchored on the entry's citeproc-generated bookmark
    # (ref-<key>) per citation, styled via the reference-doc's own
    # "Hyperlink" character style rather than an unverified Word fallback.
    # This only checks the *source* side (quartifyr's own responsibility --
    # requesting the hyperlink and styling it). Whether each hyperlink's
    # *target* bookmark actually survives reportifyr's pass-2 fill intact
    # is checked separately below, and known-flaky -- see that check's
    # comment for why this doesn't also assert it here.
    checks.append(
        (
            "both in-text citations are real hyperlinks to their bibliography entry",
            document_xml.count('<w:hyperlink w:anchor="ref-') == 2
            and document_xml.count('<w:rStyle w:val="Hyperlink"') == 2,
        )
    )

    # Whether the hyperlink *targets* above actually still resolve depends
    # on reportifyr's pass-2 fill, which is known (see the shell bookmark
    # check earlier in this file) to sometimes corrupt an unrelated
    # bookmark's close tag when its own footnote-bookmark id coincidentally
    # collides with a citeproc bookmark's id. Reported as a warning, not a
    # hard failure -- a reportifyr bug outside quartifyr's own pass-1
    # responsibility shouldn't fail this repo's own test suite.
    final_start_ids = set(re.findall(r'<w:bookmarkStart w:id="(\d+)"', document_xml))
    final_end_ids = set(re.findall(r'<w:bookmarkEnd w:id="(\d+)"', document_xml))
    unpaired_in_final = final_start_ids - final_end_ids
    if unpaired_in_final:
        print(
            f"WARN: final delivered docx has unpaired bookmark id(s) {sorted(unpaired_in_final)} "
            "-- likely reportifyr's remove_bookmarks() id-collision bug (see smoke_test.py comment); "
            "the affected citation's hyperlink may not navigate in Word",
            file=sys.stderr,
        )

    # crossref-hyperlinks: left at its default (true) in this demo -- both
    # quarto-plus's {{< crossref "TblPkSummary" >}} and this extension's
    # own {{< appendix_crossref "StatisticalMethods" >}} stay hyperlinked.
    # "false" and "same-page" are exercised by tests/python/test_layout.py
    # and tests/python/test_same_page_crossrefs.py instead of here, since
    # this demo's render.R doesn't run the separate, opt-in
    # resolve-same-page-crossrefs step "same-page" needs to mean anything
    # beyond the default -- see report.qmd's comment for why.
    instr_texts = re.findall(r"<w:instrText[^>]*>([^<]*)</w:instrText>", document_xml)
    checks.append(
        (
            "both cross-references are hyperlinked (default crossref-hyperlinks: true)",
            any(t.strip() == "REF TblPkSummary \\h" for t in instr_texts)
            and any(t.strip() == "REF StatisticalMethods \\h" for t in instr_texts),
        )
    )

    checks.append(("document split into title/front-matter/body sections", len(document.sections) == 3))
    title_section, front_matter_section, body_section = document.sections

    header_expected = "ACME-001 - RPT-2026-014\tFINAL"
    checks.append(("dynamic 2-zone page header resolved from header-format:", title_section.header.paragraphs[0].text == header_expected))
    checks.append(
        (
            "header applies to all three sections",
            title_section.header.paragraphs[0].text
            == front_matter_section.header.paragraphs[0].text
            == body_section.header.paragraphs[0].text,
        )
    )
    # Tab stop must actually be flush right, not landing on the Header
    # style's own inherited center tab -- see layout.py's
    # _clear_inherited_tab_stops() docstring for the real-Word-verified bug
    # this guards against.
    header_tabs = title_section.header.paragraphs[0].paragraph_format.tab_stops
    checks.append(("header has exactly one active (non-cleared) tab stop", sum(1 for t in header_tabs if t.alignment != WD_TAB_ALIGNMENT.CLEAR) == 1))

    confidential_expected = "Confidential — Do Not Distribute"
    checks.append(("title page footer has confidentiality label + roman page number", title_section.footer.paragraphs[0].text == f"{confidential_expected}\t1" and any("PAGE" in p._p.xml for p in title_section.footer.paragraphs)))
    checks.append(("front-matter footer has confidentiality label + roman page number", front_matter_section.footer.paragraphs[0].text == f"{confidential_expected}\t1" and any("PAGE" in p._p.xml for p in front_matter_section.footer.paragraphs)))
    checks.append(("body footer has confidentiality label + arabic page number", body_section.footer.paragraphs[0].text == f"{confidential_expected}\t1" and any("PAGE" in p._p.xml for p in body_section.footer.paragraphs)))
    footer_tabs = body_section.footer.paragraphs[0].paragraph_format.tab_stops
    checks.append(("footer has exactly one active (non-cleared) tab stop", sum(1 for t in footer_tabs if t.alignment != WD_TAB_ALIGNMENT.CLEAR) == 1))

    title_pg_num_type = title_section._sectPr.find(qn("w:pgNumType"))
    checks.append(
        (
            "title page numbers in lowercase roman starting at i",
            title_pg_num_type is not None
            and title_pg_num_type.get(qn("w:start")) == "1"
            and title_pg_num_type.get(qn("w:fmt")) == "lowerRoman",
        )
    )
    front_matter_pg_num_type = front_matter_section._sectPr.find(qn("w:pgNumType"))
    checks.append(
        (
            "rest of front matter continues the roman sequence starting at ii",
            front_matter_pg_num_type is not None
            and front_matter_pg_num_type.get(qn("w:start")) == "2"
            and front_matter_pg_num_type.get(qn("w:fmt")) == "lowerRoman",
        )
    )
    body_pg_num_type = body_section._sectPr.find(qn("w:pgNumType"))
    checks.append(
        (
            "body section numbers in arabic, restarting at 1",
            body_pg_num_type is not None
            and body_pg_num_type.get(qn("w:start")) == "1"
            and body_pg_num_type.get(qn("w:fmt")) == "decimal",
        )
    )

    failed = [name for name, ok in checks if not ok]
    for name, ok in checks:
        print(f"{'PASS' if ok else 'FAIL'}: {name}")

    if failed:
        print(f"\n{len(failed)}/{len(checks)} checks failed", file=sys.stderr)
        return 1

    print(f"\nAll {len(checks)} checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
