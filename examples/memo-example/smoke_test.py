#!/usr/bin/env python3
"""End-to-end smoke test for the quartifyr memo example.

Runs the real two-pass pipeline (Rscript render.R --final, which drives
Quarto + reportifyr exactly as a memo author would) and asserts on the
resulting docx: the memo cover (logo + left-aligned MEMORANDUM banner +
To/From/Date/Re/Cc grid) rendered correctly, the budget-timeline figure
filled in with a real caption, no leftover reportifyr magic strings, and
-- unlike examples/demo-report -- the *absence* of any title-page/
signature-page/synopsis/ToC artifacts, confirming title_page.lua/
signature_page.lua/synopsis.lua all stayed no-ops for a memo project.

Requires the full toolchain: R (with this repo's renv-managed packages),
Quarto, and the Python venv. Run from anywhere:

    python3 examples/memo-example/smoke_test.py
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import docx
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_DIR = Path(__file__).resolve().parent
REPO_ROOT = PROJECT_DIR.parent.parent


def _iter_paragraph_text(document: docx.Document):
    for p in document.element.body.iter(qn("w:p")):
        texts = p.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts)
        if text.strip():
            yield text


def main() -> int:
    if shutil.which("Rscript") is None:
        print("SKIP: Rscript not found on PATH", file=sys.stderr)
        return 0
    if shutil.which("quarto") is None:
        print("SKIP: quarto not found on PATH", file=sys.stderr)
        return 0

    # Same drift check examples/demo-report/smoke_test.py runs -- the
    # script covers every physical copy of the extension, both examples'
    # and the repo-root one (see scripts/sync_demo_extension.py), so this
    # one call validates this project's copy too, not just demo-report's.
    sync_check = subprocess.run(
        ["python3", str(REPO_ROOT / "scripts" / "sync_demo_extension.py"), "--check"],
        capture_output=True,
        text=True,
    )
    if sync_check.returncode != 0:
        print("FAIL: an examples/*/_extensions/quartifyr/ copy has drifted from the canonical one", file=sys.stderr)
        print(sync_check.stdout, file=sys.stderr)
        print(sync_check.stderr, file=sys.stderr)
        return 1

    # Same drift check examples/demo-report/smoke_test.py runs -- templates/
    # org-reference.docx is shared, committed toolkit source, not something
    # this example owns a copy of.
    template_check = subprocess.run(
        ["python3", str(REPO_ROOT / "scripts" / "check_template_freshness.py"), "--check"],
        capture_output=True,
        text=True,
    )
    if template_check.returncode != 0:
        print("FAIL: inst/templates/org-reference.docx has drifted from inst/python/styles/default.yaml", file=sys.stderr)
        print(template_check.stdout, file=sys.stderr)
        print(template_check.stderr, file=sys.stderr)
        return 1

    print("Running Rscript render.R --final ...")
    result = subprocess.run(
        ["Rscript", "render.R", "--final"],
        cwd=PROJECT_DIR,
        capture_output=True,
        text=True,
        timeout=180,
    )
    if result.returncode != 0:
        print("FAIL: render.R exited non-zero", file=sys.stderr)
        print(result.stdout, file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return 1

    final_docx = PROJECT_DIR / "report" / "final" / "report-final.docx"
    if not final_docx.exists():
        print(f"FAIL: expected output not found: {final_docx}", file=sys.stderr)
        return 1

    document = docx.Document(str(final_docx))
    all_text = list(_iter_paragraph_text(document))
    joined = " | ".join(all_text)

    checks: list[tuple[str, bool]] = []

    checks.append(("no leftover {rpfy}: magic strings", "{rpfy}:" not in joined))

    # build_template.py's reference-doc sets <w:updateFields w:val="true"/>
    # in settings.xml so Word auto-recalculates every field the moment a
    # human opens the delivered file -- confirmed here to survive the full
    # pipeline into the actual delivered docx, same as demo-report's own
    # check.
    update_fields = document.settings.element.find(qn("w:updateFields"))
    checks.append(
        (
            "delivered docx auto-recalculates fields on open in Word (w:updateFields survived the full pipeline)",
            update_fields is not None and update_fields.get(qn("w:val")) == "true",
        )
    )

    checks.append(("MEMORANDUM banner rendered", "MEMORANDUM" in joined))
    checks.append(("document-status stamp rendered", "DRAFT" in joined or "FINAL" in joined))

    with zipfile.ZipFile(final_docx) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
        images = [n for n in z.namelist() if n.startswith("word/media/")]

    # Left-aligned (no explicit <w:jc w:val="center"/>), unlike
    # title_page.lua's centered Title -- a real memo reads top-left, not
    # centered like a formal report cover. See memo_cover.lua's
    # file-header comment.
    banner_match = re.search(r'<w:pStyle w:val="Heading1"/>.*?MEMORANDUM', document_xml, re.DOTALL)
    checks.append(
        (
            "MEMORANDUM banner is left-aligned, not centered",
            banner_match is not None and 'w:jc w:val="center"' not in banner_match.group(0),
        )
    )

    checks.append(("logo image embedded on the memo cover", len(images) >= 1))
    # Style ID "LogoLeft" (no space), not the display name "Logo Left" --
    # raw OOXML pStyle references match by ID, see _extensions/quartifyr/
    # README.md's "A pStyle gotcha" section.
    checks.append(("logo defaults to left alignment ('LogoLeft' style used)", 'w:pStyle w:val="LogoLeft"' in document_xml))

    def _table_row_texts(table) -> list[list[str]]:
        return [[cell.text for cell in row.cells] for row in table.rows]

    all_rows = [row for table in document.tables for row in _table_row_texts(table)]
    # Still exactly one table: title-page-extra's "Memo Number" row (see
    # report.qmd's comment -- memo_number reused via YAML anchor/alias,
    # not retyped) is appended into the same To/From/Date/Re/Cc grid by
    # memo_cover.lua, not a second table.
    checks.append(("exactly one table present (the To/From/Date/Re/Cc + title-page-extra grid)", len(document.tables) == 1))
    checks.append(("grid has the expected labels and values, including the title-page-extra row", all_rows == [
        ["To", "Jane Doe, CFO"],
        ["From", "John Smith, Controller"],
        ["Date", "2026-08-12"],
        ["Re", "Q3 Budget Review Timeline"],
        ["Cc", "Finance Committee"],
        ["Memo Number", "MEMO-2026-014"],
    ]))

    # The budget-timeline figure (scripts/01_analysis.R -> OUTPUTS/figures/
    # budget-timeline.png), filled in by reportifyr from the {rpfy}:
    # placeholder in report.qmd, with a live SEQ-Figure caption and a
    # hyperlinked crossref back to it in the following paragraph.
    checks.append(("budget-timeline figure embedded as a real image", len(images) >= 2))
    checks.append(("figure caption rendered with live numbering", "Figure 1" in joined and "Proposed Q3 budget review timeline" in joined))
    checks.append(("figure footnote (source/notes) rendered", "Proposed Q3 budget review milestones." in joined))
    instr_texts = re.findall(r"<w:instrText[^>]*>([^<]*)</w:instrText>", document_xml)
    checks.append(("figure crossref is hyperlinked (default crossref-hyperlinks: true)", any(t.strip() == "REF FigTimeline \\h" for t in instr_texts)))

    # indent-headers: false (report.qmd) -- quarto-plus's header.lua would
    # otherwise insert a leading tab before every heading (meant to line
    # up with a "1." from number-sections, which this memo also disables).
    checks.append(("headings aren't tab-indented (indent-headers: false)", "\tPurpose" not in joined and "\tProposed Schedule" not in joined and "\tNext Steps" not in joined))

    # The whole point of this example: none of a report's other
    # front-matter pieces render for a memo project, since none of their
    # triggering frontmatter (title/contributors/approvers/synopsis) or
    # quarto-plus ToC/LoF/LoT/abbreviations divs are present in report.qmd.
    checks.append(("title_page.lua stayed a no-op (no 'Title Page' ToC marker)", "Title Page" not in joined))
    checks.append(("signature_page.lua stayed a no-op (no 'Signatures' page)", "Signatures" not in joined))
    checks.append(("synopsis.lua stayed a no-op (no 'Synopsis' section)", "Synopsis" not in joined))
    checks.append(("no 'Table of Contents' heading", "Table of Contents" not in joined))
    checks.append(("no 'List of Figures' heading", "List of Figures" not in joined))
    checks.append(("no 'List of Tables' heading", "List of Tables" not in joined))
    # A bare substring check would false-positive on reportifyr's own
    # auto-generated per-figure footnote line ("...\nAbbreviations: N/A"),
    # which has nothing to do with quarto-plus's Abbreviations *section* --
    # check for it as a standalone heading paragraph instead.
    checks.append(("no 'Abbreviations' heading", "Abbreviations" not in all_text))

    # Unlike title_page.lua, memo_cover.lua deliberately does NOT emit
    # quartifyr-front-matter-start -- giving the cover its own OOXML
    # section here would butt directly against body-start's own section
    # boundary with nothing in between, which renders as a fully
    # blank page (confirmed in practice -- see utils.
    # front_matter_start_bookmark()'s comment). Only body-start should
    # exist, giving a clean 2-section split with no blank page.
    checks.append(("no quartifyr-front-matter-start bookmark (memo cover shares a section with the rest of the front matter)", "quartifyr-front-matter-start" not in document_xml))
    checks.append(("exactly one quartifyr-body-start bookmark", document_xml.count("quartifyr-body-start") == 1))

    checks.append(("document split into front-matter (cover) and body sections only -- no blank page", len(document.sections) == 2))
    front_matter_section, body_section = document.sections

    header_expected = "ACME-001 - MEMO-2026-014\tFINAL"
    checks.append(("dynamic 2-zone page header resolved from header-format:", front_matter_section.header.paragraphs[0].text == header_expected))
    checks.append(
        (
            "header applies to both sections",
            front_matter_section.header.paragraphs[0].text == body_section.header.paragraphs[0].text,
        )
    )
    header_tabs = front_matter_section.header.paragraphs[0].paragraph_format.tab_stops
    checks.append(("header has exactly one active (non-cleared) tab stop", sum(1 for t in header_tabs if t.alignment != WD_TAB_ALIGNMENT.CLEAR) == 1))

    confidential_expected = "Internal Use Only"
    # Deliberately different from examples/demo-report's title page: the
    # cover's footer shows the confidentiality label with NO page number
    # (matching a typical memo's unnumbered cover) -- see
    # utils.front_matter_start_bookmark()'s comment for why.
    checks.append(
        (
            "cover/front-matter footer has confidentiality label and no page number",
            front_matter_section.footer.paragraphs[0].text == confidential_expected
            and not any("PAGE" in p._p.xml for p in front_matter_section.footer.paragraphs),
        )
    )
    checks.append(("body footer has confidentiality label + arabic page number", body_section.footer.paragraphs[0].text == f"{confidential_expected}\t1" and any("PAGE" in p._p.xml for p in body_section.footer.paragraphs)))

    body_pg_num_type = body_section._sectPr.find(qn("w:pgNumType"))
    checks.append(
        (
            "body section numbers in arabic, restarting at 1",
            body_pg_num_type is not None
            and body_pg_num_type.get(qn("w:start")) == "1"
            and body_pg_num_type.get(qn("w:fmt")) == "decimal",
        )
    )

    failed = [name for name, ok in checks if not ok]
    for name, ok in checks:
        print(f"{'PASS' if ok else 'FAIL'}: {name}")

    if failed:
        print(f"\n{len(failed)}/{len(checks)} checks failed", file=sys.stderr)
        return 1

    print(f"\nAll {len(checks)} checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
