#!/usr/bin/env python3
"""Unit-style test for synopsis.lua's four `synopsis-style:` options.

Renders scripts/fixtures/synopsis_style.qmd once per style
(definition-list, inline, table, false) via plain `quarto render`
against the committed reference-doc -- no R, no reportifyr, no
render_report() -- and asserts on the resulting docx's raw paragraph/
table structure (pStyle/rStyle/w:ind, not rendered pixels). This is
deliberately narrower than examples/demo-report/smoke_test.py: it
isolates synopsis.lua's own layout logic from the full two-pass
pipeline, the same way scripts/quarto_only_smoke_test.py isolates the
Quarto-only render path -- but needs python-docx (the styling/ venv) for
the structural digging the checks below do, unlike that script's
stdlib-only regex approach.

Run from anywhere (with the styling/ venv active, or via
`uv run --project styling`):

    python3 scripts/test_synopsis_styles.py
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = REPO_ROOT / "inst" / "templates" / "org-reference.docx"
FIXTURE = REPO_ROOT / "scripts" / "fixtures" / "synopsis_style.qmd"


def _render(style_metadata_arg: str | None, work_dir: Path) -> subprocess.CompletedProcess:
    args = ["quarto", "render", "synopsis_style.qmd", "--to", "docx", "--reference-doc", str(TEMPLATE)]
    if style_metadata_arg is not None:
        args += ["-M", style_metadata_arg]
    return subprocess.run(args, cwd=work_dir, capture_output=True, text=True)


def _body_paragraphs(doc):
    import docx
    from docx.oxml.ns import qn

    paragraphs = list(doc.element.body.iter(qn("w:p")))
    texts = ["".join(t.text or "" for t in p.findall(".//" + qn("w:t"))) for p in paragraphs]
    return paragraphs, texts


def _check_definition_list(work_dir: Path, checks: list[tuple[str, bool]]) -> None:
    import docx
    from docx.oxml.ns import qn

    result = _render(None, work_dir)  # no override -- proves this is really the default
    checks.append(("definition-list (default, no synopsis-style: set): render succeeds", result.returncode == 0))
    if result.returncode != 0:
        print(result.stdout, file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return

    doc = docx.Document(str(work_dir / "synopsis_style.docx"))
    paragraphs, texts = _body_paragraphs(doc)

    label_idx = next((i for i, t in enumerate(texts) if t == "Objectives"), None)
    checks.append(("definition-list: label rendered as its own paragraph", label_idx is not None))

    value_ind = None
    if label_idx is not None:
        ppr = paragraphs[label_idx + 1].find(qn("w:pPr"))
        if ppr is not None:
            value_ind = ppr.find(qn("w:ind"))
    checks.append(("definition-list: value has no w:ind override (inherits the indented 'Synopsis Value' style)", value_ind is None))

    label_pstyle = paragraphs[label_idx].find(".//" + qn("w:pStyle")) if label_idx is not None else None
    checks.append(("definition-list: label paragraph uses the 'SynopsisLabel' style", label_pstyle is not None and label_pstyle.get(qn("w:val")) == "SynopsisLabel"))


def _check_inline(work_dir: Path, checks: list[tuple[str, bool]]) -> None:
    import docx
    from docx.oxml.ns import qn

    result = _render("synopsis-style:inline", work_dir)
    checks.append(("inline: render succeeds", result.returncode == 0))
    if result.returncode != 0:
        print(result.stdout, file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return

    doc = docx.Document(str(work_dir / "synopsis_style.docx"))
    paragraphs, texts = _body_paragraphs(doc)

    single_line_idx = next((i for i, t in enumerate(texts) if t == "Objectives:  A single plain-text value line."), None)
    checks.append(("inline: single-line value merges into one 'Label:  value' paragraph", single_line_idx is not None))

    rstyle = paragraphs[single_line_idx].find(".//" + qn("w:rStyle")) if single_line_idx is not None else None
    checks.append(("inline: label run uses the 'SynopsisInlineLabel' character style", rstyle is not None and rstyle.get(qn("w:val")) == "SynopsisInlineLabel"))

    # Regression check: LibreOffice does not reliably give a run's
    # referenced character style priority over its paragraph style's own
    # bold (confirmed by rendering -- the label came out plain and the
    # value came out bold, backwards from both styles' definitions, even
    # though "SynopsisInlineLabel" should win per the OOXML spec). Direct
    # run-level <w:b/> is unambiguous in both Word and LibreOffice, so
    # both runs must carry an explicit one rather than relying on either
    # style's own rPr cascading correctly.
    runs = paragraphs[single_line_idx].findall(qn("w:r")) if single_line_idx is not None else []
    label_run_b = runs[0].find(qn("w:rPr") + "/" + qn("w:b")) if len(runs) > 0 else None
    value_run_b = runs[1].find(qn("w:rPr") + "/" + qn("w:b")) if len(runs) > 1 else None
    checks.append(("inline: label run has an explicit direct <w:b/>, not just the rStyle reference", label_run_b is not None and label_run_b.get(qn("w:val")) != "0"))
    checks.append(("inline: value run has an explicit direct <w:b w:val=\"0\"/> overriding the paragraph style's bold", value_run_b is not None and value_run_b.get(qn("w:val")) == "0"))

    # Results' first value line is plain text ("First line of results."),
    # so the label still merges into it even though a second (image) line
    # follows -- only the *first* line's shape decides mergeability.
    results_merge_idx = next((i for i, t in enumerate(texts) if t == "Results:  First line of results."), None)
    checks.append(("inline: label merges into the first line even when more lines (an image) follow (Results)", results_merge_idx is not None))

    # Figure's first value line IS an image -- can't run "Label:  " into
    # a picture, so the label stands alone, but still bold/larger/colon
    # via inline_label_only_paragraph, not plain_label_paragraph's
    # unstyled look; its second line ("Caption after the image.") still
    # follows underneath, flush -- indentation is a definition-list
    # trait only, inline never indents.
    figure_label_idx = next((i for i, t in enumerate(texts) if t == "Figure:"), None)
    checks.append(("inline: a row whose first line is an image falls back to a standalone label (Figure)", figure_label_idx is not None))

    figure_rstyle = paragraphs[figure_label_idx].find(".//" + qn("w:rStyle")) if figure_label_idx is not None else None
    checks.append(("inline: standalone fallback label still uses the bold/larger 'SynopsisInlineLabel' style", figure_rstyle is not None and figure_rstyle.get(qn("w:val")) == "SynopsisInlineLabel"))

    caption_idx = next((i for i, t in enumerate(texts) if t == "Caption after the image."), None)
    caption_ind = None
    if caption_idx is not None:
        ppr = paragraphs[caption_idx].find(qn("w:pPr"))
        if ppr is not None:
            caption_ind = ppr.find(qn("w:ind"))
    checks.append(
        (
            "inline: a line after the fallback label is flush (w:ind left=0 override present, not inherited indent)",
            caption_ind is not None and caption_ind.get(qn("w:left")) == "0",
        )
    )


def _check_table(work_dir: Path, checks: list[tuple[str, bool]]) -> None:
    import docx
    from docx.oxml.ns import qn

    result = _render("synopsis-style:table", work_dir)
    checks.append(("table: render succeeds", result.returncode == 0))
    if result.returncode != 0:
        print(result.stdout, file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return

    output = result.stdout + result.stderr
    checks.append(
        (
            "table: quarto.log.warning() fires about reportifyr's footnote placement",
            "reportifyr" in output and "footnote" in output and "table cell" in output,
        )
    )

    doc = docx.Document(str(work_dir / "synopsis_style.docx"))
    checks.append(("table: a real w:tbl is present", len(doc.tables) >= 1))

    table_style = None
    row_texts = []
    if doc.tables:
        tbl = doc.tables[0]
        table_style = tbl._element.find(qn("w:tblPr") + "/" + qn("w:tblStyle"))
        row_texts = [(row.cells[0].text, row.cells[1].text) for row in tbl.rows]
    checks.append(("table: uses the bordered 'TableGrid' style, not the borderless title-page one", table_style is not None and table_style.get(qn("w:val")) == "TableGrid"))
    checks.append(("table: Objectives row present with its value in the second cell", ("Objectives", "A single plain-text value line.") in row_texts))

    # Regression check: each value line must be its own <w:p> in the
    # cell, not joined into one paragraph via <w:br/>. Confirmed by
    # rendering a row with an embedded image through the full pipeline
    # that br-joining is destructive here -- reportifyr's
    # remove_magic_strings() deletes a cell paragraph outright once it
    # contains the magic string but not the (separately inserted)
    # picture, taking every other br-joined line in that same paragraph
    # down with it. This check can't exercise reportifyr itself (this
    # script deliberately has no R dependency), but it can confirm the
    # per-line-paragraph shape that keeps that deletion scoped to just
    # the magic-string line.
    results_value_cell = None
    for row in doc.tables[0].rows if doc.tables else []:
        if row.cells[0].text == "Results":
            results_value_cell = row.cells[1]
            break
    results_value_paragraph_count = len(results_value_cell._tc.findall(qn("w:p"))) if results_value_cell is not None else 0
    checks.append(
        (
            "table: a multi-line row's value cell has one <w:p> per line, not one <w:br/>-joined paragraph",
            results_value_paragraph_count >= 2,  # Results has 1 text line + 1 image line in the fixture
        )
    )


def _check_false(work_dir: Path, checks: list[tuple[str, bool]]) -> None:
    import docx
    from docx.oxml.ns import qn

    result = _render("synopsis-style:false", work_dir)
    checks.append(("false: render succeeds", result.returncode == 0))
    if result.returncode != 0:
        print(result.stdout, file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return

    doc = docx.Document(str(work_dir / "synopsis_style.docx"))
    _, texts = _body_paragraphs(doc)
    checks.append(("false: heading still renders", "Synopsis" in texts))
    checks.append(("false: no synopsis content rendered at all", "Objectives" not in texts and "Results" not in texts))
    checks.append(("false: no leftover {rpfy}: magic strings either (nothing rendered means nothing to leak)", not any("{rpfy}:" in t for t in texts)))


def main() -> int:
    if shutil.which("quarto") is None:
        print("SKIP: quarto not found on PATH", file=sys.stderr)
        return 0
    try:
        import docx  # noqa: F401
    except ImportError:
        print("SKIP: python-docx not importable -- activate styling/'s venv", file=sys.stderr)
        return 0
    if not TEMPLATE.exists():
        print(f"FAIL: {TEMPLATE} not found -- it should be committed to the repo", file=sys.stderr)
        return 1

    checks: list[tuple[str, bool]] = []
    with tempfile.TemporaryDirectory(prefix="quartifyr-synopsis-styles-") as tmp:
        work_dir = Path(tmp)
        shutil.copytree(REPO_ROOT / "inst" / "extensions" / "quartifyr", work_dir / "_extensions" / "quartifyr")
        shutil.copy(FIXTURE, work_dir / "synopsis_style.qmd")

        _check_definition_list(work_dir, checks)
        _check_inline(work_dir, checks)
        _check_table(work_dir, checks)
        _check_false(work_dir, checks)

    failed = [n for n, ok in checks if not ok]
    for n, ok in checks:
        print(f"{'PASS' if ok else 'FAIL'}: {n}")

    if failed:
        print(f"\n{len(failed)}/{len(checks)} checks failed", file=sys.stderr)
        return 1

    print(f"\nAll {len(checks)} checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
