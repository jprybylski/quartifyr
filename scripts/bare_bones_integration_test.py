#!/usr/bin/env python3
"""Integration test for the "Using the pieces directly" path documented
in the repo-root README: plain `quarto render` + `quartifyr-styling
apply-layout` + a direct `reportifyr::build_report()` call, with none of
`r/`'s `render_report()` orchestration, its `rv`-managed environment, or
its `report/shell`/`report/draft`/`report/final` directory convention
involved.

This exists to prove that documented claim is actually true, not just
written down: the shell renders from a temp project containing nothing
but a physical copy of the extensions and a bare `report.qmd` -- no
`_quarto.yml`, no `render.R`, no `r/` at all.

Reuses `examples/demo-report`'s already-initialized reportifyr project
(its `OUTPUTS/`, `standard_footnotes.yaml`, `config.yaml`, and rv-managed
`reportifyr`/`pyro` installation) as the *reportifyr* side of the test --
having a working reportifyr project on disk somewhere is a reportifyr
prerequisite, unrelated to what's being tested here -- while proving
`docx_in`/`docx_out` can point anywhere, with no shared directory
structure required between the Quarto render and the reportifyr fill.

Requires the full toolchain: Quarto, R (with examples/demo-report's
rv-managed packages already synced and initialized), and the styling/
venv. Run from anywhere:

    python3 scripts/bare_bones_integration_test.py
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import docx
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
DEMO_DIR = REPO_ROOT / "examples" / "demo-report"
REFERENCE_DOC = REPO_ROOT / "templates" / "org-reference.docx"
BUILD_REPORT_HELPER = Path(__file__).resolve().parent / "_bare_bones_build_report.R"

REPORT_QMD = """---
title: "Bare-Bones Pieces Integration Test"
document-status: "draft"
confidentiality: "Confidential — Do Not Distribute"
project: "ACME-001"
report_number: "RPT-BARE-001"
header-format: "{project} - {report_number}"
filters:
  - quarto-plus
  - quartifyr
toc-style-map:
  - style: "Title"
    level: 1
---

::: .toc
:::

{{< pagebreak >}}

{{< body-start >}}

# Introduction

Proves the pieces work independently of `render_report()`.

{{< tbl_caption "Table 1" "Reused demo PK summary" >}}

{rpfy}:pk-summary.csv

Table content is filled by a direct `reportifyr::build_report()` call.
"""


def _iter_paragraph_text(document: docx.Document):
    for p in document.element.body.iter(qn("w:p")):
        texts = p.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts)
        if text.strip():
            yield text


def main() -> int:
    if shutil.which("quarto") is None:
        print("SKIP: quarto not found on PATH", file=sys.stderr)
        return 0
    if shutil.which("Rscript") is None:
        print("SKIP: Rscript not found on PATH", file=sys.stderr)
        return 0
    quartifyr_styling_bin = REPO_ROOT / ".venv" / "bin" / "quartifyr-styling"
    if not quartifyr_styling_bin.exists():
        print(f"SKIP: quartifyr-styling not found at {quartifyr_styling_bin}", file=sys.stderr)
        return 0
    if not REFERENCE_DOC.exists():
        print(f"SKIP: reference-doc not found at {REFERENCE_DOC} -- run quartifyr-styling build first", file=sys.stderr)
        return 0
    if not (DEMO_DIR / ".report_init.json").exists():
        print(
            f"SKIP: {DEMO_DIR} isn't an initialized reportifyr project (no .report_init.json) -- "
            "run Rscript -e 'reportifyr::initialize_report_project(project_dir = getwd())' there first",
            file=sys.stderr,
        )
        return 0

    with tempfile.TemporaryDirectory(prefix="quartifyr-bare-bones-") as tmp:
        project_dir = Path(tmp)

        # Deliberately NOT a render_report()-style project: no
        # _quarto.yml, no report/shell//report/draft//report/final
        # layout. Just the extensions, physically copied (matching what
        # `quarto add` would leave behind), and a shell .qmd.
        shutil.copytree(REPO_ROOT / "_extensions" / "quartifyr", project_dir / "_extensions" / "quartifyr")
        # quarto-plus isn't vendored at the repo root (quartifyr is the
        # only canonical extension there) -- reuse the copy `quarto add
        # A2-ai/quarto-plus` already left in the demo project.
        shutil.copytree(DEMO_DIR / "_extensions" / "A2-ai", project_dir / "_extensions" / "A2-ai")
        qmd_path = project_dir / "report.qmd"
        qmd_path.write_text(REPORT_QMD)

        shell_docx = project_dir / "shell.docx"
        print("1/3: quarto render (no _quarto.yml, no output-dir convention)...")
        result = subprocess.run(
            [
                "quarto", "render", qmd_path.name,
                "--to", "docx",
                "--reference-doc", str(REFERENCE_DOC),
                "--output", shell_docx.name,
                "-M", "document-status:DRAFT",
            ],
            cwd=project_dir,
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            print("FAIL: quarto render failed", file=sys.stderr)
            print(result.stdout, file=sys.stderr)
            print(result.stderr, file=sys.stderr)
            return 1
        if not shell_docx.exists():
            print(f"FAIL: expected {shell_docx} to exist after quarto render", file=sys.stderr)
            return 1

        print("2/3: quartifyr-styling apply-layout...")
        result = subprocess.run(
            [
                str(quartifyr_styling_bin), "apply-layout",
                "--docx", str(shell_docx),
                "--qmd", str(qmd_path),
                "--status", "draft",
            ],
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            print("FAIL: quartifyr-styling apply-layout failed", file=sys.stderr)
            print(result.stdout, file=sys.stderr)
            print(result.stderr, file=sys.stderr)
            return 1

        filled_docx = project_dir / "filled.docx"
        final_docx = project_dir / "final.docx"
        print("3/3: reportifyr::build_report() + finalize_document() called directly (no make_doc_dirs()/render_report())...")
        result = subprocess.run(
            ["Rscript", str(BUILD_REPORT_HELPER), str(shell_docx), str(filled_docx), str(final_docx)],
            cwd=DEMO_DIR,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            print("FAIL: reportifyr::build_report()/finalize_document() failed", file=sys.stderr)
            print(result.stdout, file=sys.stderr)
            print(result.stderr, file=sys.stderr)
            return 1
        if not final_docx.exists():
            print(f"FAIL: expected {final_docx} to exist after finalize_document()", file=sys.stderr)
            return 1

        document = docx.Document(str(final_docx))
        joined = " | ".join(_iter_paragraph_text(document))

        checks: list[tuple[str, bool]] = []
        checks.append(("no leftover {rpfy}: magic strings", "{rpfy}:" not in joined))
        checks.append(("title rendered", "Bare-Bones Pieces Integration Test" in joined))
        checks.append(("apply-layout split the document into sections", len(document.sections) >= 2))

        header_text = document.sections[0].header.paragraphs[0].text if document.sections[0].header.paragraphs else ""
        checks.append(
            (
                "header shows project/report_number/status (apply-layout ran)",
                "ACME-001 - RPT-BARE-001" in header_text and "DRAFT" in header_text,
            )
        )

        all_rows = [[cell.text for cell in row.cells] for table in document.tables for row in table.rows]
        checks.append(
            ("PK summary table filled with real data from the reused demo OUTPUTS/", any({"Cmax", "Cmin"}.issubset(set(row)) for row in all_rows))
        )

        failed = [name for name, ok in checks if not ok]
        for name, ok in checks:
            print(f"{'PASS' if ok else 'FAIL'}: {name}")

        if failed:
            print(f"\n{len(failed)}/{len(checks)} checks failed", file=sys.stderr)
            return 1

        print(f"\nAll {len(checks)} checks passed.")
        return 0


if __name__ == "__main__":
    raise SystemExit(main())
