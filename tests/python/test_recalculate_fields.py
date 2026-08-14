import os
import shutil
from pathlib import Path

import docx
import pytest
from docx.oxml.ns import qn

from quartifyr_styling.recalculate_fields import (
    FieldRecalculationError,
    _seed_profile,
    recalculate_fields,
)

HAS_SOFFICE = shutil.which("soffice") is not None
# The real LibreOffice integration test is slow (tens of seconds) and has
# been observed to hang intermittently in some sandboxed/CI-like
# environments for reasons not yet root-caused (see recalculate_fields.py's
# module docstring and the repo-root README.md) -- opt in explicitly rather than
# letting it silently slow down or destabilize the default test run.
RUN_SLOW_LO_TEST = os.environ.get("QUARTIFYR_RUN_SLOW_TESTS") == "1"


def test_missing_docx_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        recalculate_fields(tmp_path / "does-not-exist.docx")


def test_missing_soffice_binary_raises(tmp_path):
    docx_path = tmp_path / "fake.docx"
    docx_path.write_bytes(b"not a real docx, just needs to exist")
    with pytest.raises(FieldRecalculationError, match="not found on PATH"):
        recalculate_fields(docx_path, soffice_bin="definitely-not-a-real-binary")


def test_seed_profile_writes_expected_macro_files(tmp_path):
    _seed_profile(tmp_path)
    standard_dir = tmp_path / "user" / "basic" / "Standard"
    assert (standard_dir / "Module1.xba").exists()
    assert (standard_dir / "script.xlb").exists()
    assert (tmp_path / "user" / "basic" / "script.xlc").exists()

    macro_text = (standard_dir / "Module1.xba").read_text()
    assert "UpdateAndSave" in macro_text
    assert "getDocumentIndexes" in macro_text


def _build_minimal_toc_docx(path: Path) -> None:
    """Build a minimal docx with a real heading and a native Word TOC field,
    mirroring what quarto-plus's table_of_contents.lua + Quarto's own
    heading output produce -- without depending on Quarto being installed.
    """
    document = docx.Document()

    toc_ooxml = """
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>
      <w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>Right-click to update field</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
    """
    # Insert raw TOC field XML directly into the body, replacing a
    # placeholder paragraph.
    from lxml import etree

    placeholder = document.add_paragraph()
    parsed = etree.fromstring(toc_ooxml.encode())
    placeholder._p.addnext(parsed)
    placeholder._p.getparent().remove(placeholder._p)

    document.add_page_break()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body text.")
    document.save(str(path))


@pytest.mark.skipif(not HAS_SOFFICE, reason="LibreOffice (soffice) not installed")
@pytest.mark.skipif(
    not RUN_SLOW_LO_TEST,
    reason="slow/flaky LibreOffice integration test; set QUARTIFYR_RUN_SLOW_TESTS=1 to run it",
)
def test_recalculate_fields_populates_toc(tmp_path):
    docx_path = tmp_path / "toc-test.docx"
    _build_minimal_toc_docx(docx_path)

    result = recalculate_fields(docx_path, timeout_seconds=90)
    assert result == docx_path

    updated = docx.Document(str(docx_path))
    all_text = []
    for p in updated.element.body.iter(qn("w:p")):
        texts = p.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts)
        if text.strip():
            all_text.append(text)

    joined = " | ".join(all_text)
    assert "Right-click to update field" not in joined
    # The recalculated ToC should now contain an entry for the heading.
    assert any("Introduction" in t for t in all_text)
