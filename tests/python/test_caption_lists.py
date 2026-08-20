from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from quartifyr_styling._caption_lists import (
    LIST_OF_FIGURES_BOOKMARK,
    LIST_OF_TABLES_BOOKMARK,
    build_caption_lists,
    collect_captions,
)

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _add_bookmark(paragraph, *, bookmark_id: str, name: str) -> None:
    for el_xml in [
        f'<w:bookmarkStart xmlns:w="{W_NS[1:-1]}" w:id="{bookmark_id}" w:name="{name}"/>',
        f'<w:bookmarkEnd xmlns:w="{W_NS[1:-1]}" w:id="{bookmark_id}"/>',
    ]:
        paragraph._p.append(etree.fromstring(el_xml))


def _add_caption_paragraph(document, *, bookmark: str, seq_field: str, local_number: str, caption_text: str):
    """Builds a paragraph matching the exact bookmarkStart+SEQ...ARABIC+tab+
    caption shape both `crossref.lua` (quarto-plus's continuous
    fig_caption/tbl_caption) and `appendix.lua` (this extension's six
    scoped caption shortcodes) emit -- see `_caption_lists.py`'s own
    `_match_caption_paragraph()`, which scans for that exact shape.
    """
    p = document.add_paragraph()
    _add_bookmark(p, bookmark_id="1", name=bookmark)

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_field} \\* ARABIC "
    run._r.append(instr)

    run = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    p.add_run(local_number)

    run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    run = p.add_run()
    run._r.append(OxmlElement("w:tab"))
    p.add_run(caption_text)
    return p


def _add_marker_bookmark(document, name: str):
    p = document.add_paragraph()
    _add_bookmark(p, bookmark_id="800003", name=name)
    return p


def test_collect_captions_finds_continuous_and_scoped_in_document_order():
    document = docx.Document()
    _add_caption_paragraph(document, bookmark="TblPkSummary", seq_field="Table", local_number="1", caption_text="PK summary")
    _add_caption_paragraph(document, bookmark="FigConcTime", seq_field="Figure", local_number="1", caption_text="Conc-time")
    _add_caption_paragraph(
        document, bookmark="FigAppendixExample", seq_field="AppendixFigure", local_number="1", caption_text="Appendix fig"
    )
    _add_caption_paragraph(
        document, bookmark="TblAppendixExample", seq_field="AppendixTable", local_number="1", caption_text="Appendix tbl"
    )
    _add_caption_paragraph(
        document, bookmark="FigSectionExample", seq_field="SectionFigure", local_number="1", caption_text="Section fig"
    )

    figures, tables = collect_captions(document)

    assert [f["bookmark"] for f in figures] == ["FigConcTime", "FigAppendixExample", "FigSectionExample"]
    assert [t["bookmark"] for t in tables] == ["TblPkSummary", "TblAppendixExample"]
    assert figures[0]["caption_text"] == "Conc-time"
    assert figures[0]["label"] == "Figure"
    assert tables[0]["caption_text"] == "PK summary"


def test_collect_captions_ignores_appendix_heading_bookmark():
    # The `appendix` heading's own SEQ Appendix \* ALPHABETIC bookmark
    # must never be mistaken for a caption -- no ARABIC switch, and even
    # under appendix-numbering: arabic its field name "Appendix" contains
    # neither "Figure" nor "Table".
    document = docx.Document()
    _add_caption_paragraph(document, bookmark="StatisticalMethods", seq_field="Appendix", local_number="A", caption_text=": Statistical Methods")

    figures, tables = collect_captions(document)
    assert figures == []
    assert tables == []


def test_collect_captions_ignores_appendix_heading_with_arabic_switch():
    document = docx.Document()
    p = document.add_paragraph()
    _add_bookmark(p, bookmark_id="1", name="StatisticalMethods")
    run = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " SEQ Appendix \\* ARABIC "
    run._r.append(instr)

    figures, tables = collect_captions(document)
    assert figures == []
    assert tables == []


def test_build_caption_lists_noop_without_marker_bookmarks(tmp_path: Path):
    document = docx.Document()
    _add_caption_paragraph(document, bookmark="FigConcTime", seq_field="Figure", local_number="1", caption_text="Conc-time")

    build_caption_lists(document)  # should not raise

    docx_path = tmp_path / "test.docx"
    document.save(str(docx_path))
    result = docx.Document(str(docx_path))
    assert "PAGEREF" not in result.element.body.xml


def test_build_caption_lists_fills_marker_with_entries_in_order(tmp_path: Path):
    document = docx.Document()
    _add_marker_bookmark(document, LIST_OF_FIGURES_BOOKMARK)
    _add_marker_bookmark(document, LIST_OF_TABLES_BOOKMARK)
    _add_caption_paragraph(document, bookmark="TblPkSummary", seq_field="Table", local_number="1", caption_text="PK summary")
    _add_caption_paragraph(document, bookmark="FigConcTime", seq_field="Figure", local_number="1", caption_text="Conc-time")
    _add_caption_paragraph(
        document, bookmark="FigAppendixExample", seq_field="AppendixFigure", local_number="1", caption_text="Appendix fig"
    )

    build_caption_lists(document)

    docx_path = tmp_path / "test.docx"
    document.save(str(docx_path))
    result = docx.Document(str(docx_path))
    body_xml = result.element.body.xml

    # Both entries anchored on the caption's own bookmark, in document
    # order, each carrying a REF (number) and PAGEREF (page) field.
    fig_idx = body_xml.index('w:anchor="FigConcTime"')
    appendix_fig_idx = body_xml.index('w:anchor="FigAppendixExample"')
    assert fig_idx < appendix_fig_idx
    assert "REF FigConcTime \\h" in body_xml
    assert "PAGEREF FigConcTime \\h" in body_xml
    assert "REF FigAppendixExample \\h" in body_xml
    assert "PAGEREF FigAppendixExample \\h" in body_xml

    tbl_idx = body_xml.index('w:anchor="TblPkSummary"')
    assert "REF TblPkSummary \\h" in body_xml
    assert "PAGEREF TblPkSummary \\h" in body_xml
    assert tbl_idx > 0

    assert body_xml.count('w:val="TOC1"') == 3
    assert 'w:val="Hyperlink"' in body_xml


def test_build_caption_lists_empty_list_leaves_marker_untouched(tmp_path: Path):
    document = docx.Document()
    _add_marker_bookmark(document, LIST_OF_FIGURES_BOOKMARK)
    # No table captions at all -- the figures marker above has none either.

    build_caption_lists(document)

    docx_path = tmp_path / "test.docx"
    document.save(str(docx_path))
    result = docx.Document(str(docx_path))
    assert "PAGEREF" not in result.element.body.xml
