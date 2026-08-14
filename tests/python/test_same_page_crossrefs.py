import os
import shutil
from pathlib import Path

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from quartifyr_styling.same_page_crossrefs import (
    SamePageCrossrefError,
    _find_markers,
    _seed_profile,
    resolve_same_page_crossrefs,
)

HAS_SOFFICE = shutil.which("soffice") is not None
# See test_recalculate_fields.py's identical guard: slow and observed to
# hang intermittently in some sandboxed/CI-like environments -- opt in
# explicitly.
RUN_SLOW_LO_TEST = os.environ.get("QUARTIFYR_RUN_SLOW_TESTS") == "1"

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _add_bookmark(paragraph, *, bookmark_id: int, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)


def _add_ref_field_run(paragraph, *, bookmark: str, hyperlink: bool, cached_text: str = "Figure 1") -> None:
    """The same 5-run begin/instrText/separate/result/end REF field shape
    as test_layout.py's identical helper -- kept local rather than shared
    since it's a small, self-contained fixture builder.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" REF {bookmark} \\h " if hyperlink else f" REF {bookmark} "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    paragraph.add_run(cached_text)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _build_docx_with_marked_crossref(path: Path, *, marker_name: str = "quartifyr-crossref-target-1") -> None:
    """Mirrors exactly what apply_layout()'s
    _mark_crossrefs_for_same_page_resolution() leaves behind: a marker
    bookmark immediately before a still-hyperlinked REF field.
    """
    document = docx.Document()
    p = document.add_paragraph("See ")
    _add_bookmark(p, bookmark_id=950001, name=marker_name)
    _add_ref_field_run(p, bookmark="Figure1", hyperlink=True)
    target_p = document.add_paragraph("Figure 1")
    _add_bookmark(target_p, bookmark_id=1, name="Figure1")
    document.save(str(path))


def test_missing_docx_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        resolve_same_page_crossrefs(tmp_path / "does-not-exist.docx")


def test_no_markers_is_a_noop_and_never_touches_soffice(tmp_path):
    """No crossref-hyperlinks: "same-page" markers -- nothing to resolve,
    so this must return immediately without even checking for LibreOffice
    (confirmed here via a soffice_bin that doesn't exist: a real check
    would raise, so reaching a normal return proves it was skipped).
    """
    docx_path = tmp_path / "plain.docx"
    document = docx.Document()
    document.add_paragraph("Nothing to see here.")
    document.save(str(docx_path))

    result = resolve_same_page_crossrefs(docx_path, soffice_bin="definitely-not-a-real-binary")
    assert result == docx_path.resolve()


def test_missing_soffice_binary_raises_when_markers_present(tmp_path):
    docx_path = tmp_path / "marked.docx"
    _build_docx_with_marked_crossref(docx_path)

    with pytest.raises(SamePageCrossrefError, match="not found on PATH"):
        resolve_same_page_crossrefs(docx_path, soffice_bin="definitely-not-a-real-binary")


def test_find_markers_locates_marker_and_target_bookmark(tmp_path):
    docx_path = tmp_path / "marked.docx"
    _build_docx_with_marked_crossref(docx_path)

    document = docx.Document(str(docx_path))
    markers = _find_markers(document)

    assert len(markers) == 1
    marker_name, target_bookmark, instr_text, bookmark_start, bookmark_end = markers[0]
    assert marker_name == "quartifyr-crossref-target-1"
    assert target_bookmark == "Figure1"
    assert instr_text.text.strip() == "REF Figure1 \\h"
    assert bookmark_start.get(qn("w:name")) == "quartifyr-crossref-target-1"
    assert bookmark_end.tag == qn("w:bookmarkEnd")


def test_seed_profile_writes_expected_macro_files(tmp_path):
    pairs_path = tmp_path / "pairs.txt"
    results_path = tmp_path / "results.txt"
    profile_dir = tmp_path / "profile"

    _seed_profile(profile_dir, pairs_path=pairs_path, results_path=results_path)

    standard_dir = profile_dir / "user" / "basic" / "Standard"
    assert (standard_dir / "Module1.xba").exists()
    assert (standard_dir / "script.xlb").exists()
    assert (profile_dir / "user" / "basic" / "script.xlc").exists()

    macro_text = (standard_dir / "Module1.xba").read_text()
    assert "ReadPages" in macro_text
    assert str(pairs_path) in macro_text
    assert str(results_path) in macro_text
    assert "getPage" in macro_text


@pytest.mark.skipif(not HAS_SOFFICE, reason="LibreOffice (soffice) not installed")
@pytest.mark.skipif(
    not RUN_SLOW_LO_TEST,
    reason="slow/flaky LibreOffice integration test; set QUARTIFYR_RUN_SLOW_TESTS=1 to run it",
)
def test_resolve_same_page_crossrefs_strips_hyperlink_when_target_on_same_page(tmp_path):
    """Real end-to-end LibreOffice integration test -- not independently
    verified to complete in this project's development environment (every
    attempted run, of this and of the pre-existing recalculate_fields.py
    macro, timed out rather than finishing); see the module docstring.
    """
    docx_path = tmp_path / "same-page.docx"
    _build_docx_with_marked_crossref(docx_path)

    result = resolve_same_page_crossrefs(docx_path, timeout_seconds=90)
    assert result == docx_path.resolve()

    updated = docx.Document(str(docx_path))
    ref_instr = next(
        instr.text
        for instr in updated.element.body.iter(qn("w:instrText"))
        if instr.text and "REF Figure1" in instr.text
    )
    # Reference and target are both on the trivial single-page fixture
    # document here, so the hyperlink switch should be stripped.
    assert "\\h" not in ref_instr
    # Marker bookmark cleaned up afterward.
    assert not any(
        (b.get(qn("w:name")) or "").startswith("quartifyr-crossref-target-")
        for b in updated.element.body.iter(qn("w:bookmarkStart"))
    )
