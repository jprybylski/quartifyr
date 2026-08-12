from pathlib import Path

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from quartifyr_styling.layout import (
    LayoutError,
    apply_layout,
    apply_layout_from_qmd,
    read_qmd_frontmatter,
    resolve_confidential_label,
    resolve_header_left_text,
)


def test_read_qmd_frontmatter_parses_yaml_block(tmp_path):
    qmd = tmp_path / "report.qmd"
    qmd.write_text("---\ntitle: Foo\nproject: ACME-001\n---\n\nBody.\n")
    fm = read_qmd_frontmatter(qmd)
    assert fm == {"title": "Foo", "project": "ACME-001"}


def test_read_qmd_frontmatter_no_frontmatter_returns_empty(tmp_path):
    qmd = tmp_path / "report.qmd"
    qmd.write_text("Just body, no frontmatter.\n")
    assert read_qmd_frontmatter(qmd) == {}


def test_resolve_header_left_text_none_when_no_header_format():
    assert resolve_header_left_text({}, "draft") is None


def test_resolve_header_left_text_substitutes_placeholders():
    fm = {"header-format": "{project} - {report_number}", "project": "ACME-001", "report_number": "RPT-1001"}
    assert resolve_header_left_text(fm, "draft") == "ACME-001 - RPT-1001"


def test_resolve_header_left_text_does_not_auto_append_status():
    # Status has its own always-shown zone on the right (see apply_layout),
    # so the left text is used verbatim -- no more auto-appending "—
    # DRAFT"/"— FINAL" the way a single-zone header used to.
    fm = {"header-format": "{project}", "project": "ACME-001"}
    assert resolve_header_left_text(fm, "draft") == "ACME-001"


def test_resolve_header_left_text_status_placeholder_available():
    fm = {"header-format": "{status}: {project}", "project": "ACME-001"}
    assert resolve_header_left_text(fm, "final") == "FINAL: ACME-001"


def test_resolve_header_left_text_missing_placeholder_raises_clear_error():
    fm = {"header-format": "{project} - {missing_field}", "project": "ACME-001"}
    with pytest.raises(LayoutError, match="missing_field"):
        resolve_header_left_text(fm, "draft")


def test_resolve_confidential_label_uses_confidentiality_field():
    fm = {"confidentiality": "Confidential — Do Not Distribute"}
    assert resolve_confidential_label(fm) == "Confidential — Do Not Distribute"


def test_resolve_confidential_label_blank_when_unset():
    assert resolve_confidential_label({}) == ""


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _add_page_field_run(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_ref_field_run(paragraph, *, bookmark: str, hyperlink: bool, cached_text: str = "Figure 1") -> None:
    """Mirrors the exact 5-run ``begin/instrText/separate/result/end``
    ``REF <bookmark> \\h`` field group quarto-plus's ``crossref`` shortcode
    and this extension's own ``appendix_crossref`` both emit (as distinct
    sibling ``<w:r>`` elements, not merged into one -- load-bearing for
    ``_match_ref_field_run_group()``, which scans for that exact shape).
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" REF {bookmark} \\h " if hyperlink else f" REF {bookmark} "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    paragraph.add_run(cached_text)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _add_bookmark(paragraph, *, bookmark_id: str, name: str) -> None:
    for el_xml in [
        f'<w:bookmarkStart xmlns:w="{W_NS[1:-1]}" w:id="{bookmark_id}" w:name="{name}"/>',
        f'<w:bookmarkEnd xmlns:w="{W_NS[1:-1]}" w:id="{bookmark_id}"/>',
    ]:
        paragraph._p.append(etree.fromstring(el_xml))


def _build_docx_with_bookmarks(path: Path, *, include_front_matter_bookmark: bool) -> None:
    """A minimal docx with a page-number footer (mirroring what
    styling/build_template.py's reference-doc produces) and
    quartifyr-body-start (and, unless disabled, quartifyr-front-matter-
    start) bookmarks, to exercise apply_layout() without needing a real
    Quarto render.
    """
    document = docx.Document()
    document.add_paragraph("Title page para")

    section = document.sections[0]
    section.footer.is_linked_to_previous = False
    _add_page_field_run(section.footer.paragraphs[0])

    if include_front_matter_bookmark:
        front_matter_p = document.add_paragraph()
        _add_bookmark(front_matter_p, bookmark_id="800002", name="quartifyr-front-matter-start")

    document.add_paragraph("Front matter para")

    body_start_p = document.add_paragraph()
    _add_bookmark(body_start_p, bookmark_id="800001", name="quartifyr-body-start")

    document.add_paragraph("Body para")
    document.save(str(path))


def test_apply_layout_three_sections_with_title_page_bookmark(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_bookmarks(docx_path, include_front_matter_bookmark=True)

    apply_layout(
        docx_path,
        header_left_text="ACME-001 - RPT-1001",
        status="draft",
        confidential_label="Confidential",
    )

    result = docx.Document(str(docx_path))
    assert len(result.sections) == 3
    title, front_matter, body = result.sections

    for section in (title, front_matter, body):
        assert section.header.paragraphs[0].text == "ACME-001 - RPT-1001\tDRAFT"

    # Title page: confidential label + roman numeral, starting at "i".
    assert title.footer.paragraphs[0].text == "Confidential\t1"
    assert "PAGE" in title.footer.paragraphs[0]._p.xml
    title_pg_num_type = title._sectPr.find(qn("w:pgNumType"))
    assert title_pg_num_type is not None
    assert title_pg_num_type.get(qn("w:start")) == "1"
    assert title_pg_num_type.get(qn("w:fmt")) == "lowerRoman"

    # Rest of front matter: same roman sequence, continuing at "ii".
    assert front_matter.footer.paragraphs[0].text == "Confidential\t1"
    assert "PAGE" in front_matter.footer.paragraphs[0]._p.xml
    front_matter_pg_num_type = front_matter._sectPr.find(qn("w:pgNumType"))
    assert front_matter_pg_num_type is not None
    assert front_matter_pg_num_type.get(qn("w:start")) == "2"
    assert front_matter_pg_num_type.get(qn("w:fmt")) == "lowerRoman"

    # Body: confidential label + arabic, restarting at 1.
    assert body.footer.paragraphs[0].text == "Confidential\t1"
    assert "PAGE" in body.footer.paragraphs[0]._p.xml
    body_pg_num_type = body._sectPr.find(qn("w:pgNumType"))
    assert body_pg_num_type is not None
    assert body_pg_num_type.get(qn("w:start")) == "1"
    assert body_pg_num_type.get(qn("w:fmt")) == "decimal"


def test_apply_layout_two_sections_without_title_page_bookmark(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_bookmarks(docx_path, include_front_matter_bookmark=False)

    apply_layout(docx_path, header_left_text="ACME-001 - RPT-1001", status="draft")

    result = docx.Document(str(docx_path))
    assert len(result.sections) == 2
    front_matter, body = result.sections

    assert front_matter.header.paragraphs[0].text == "ACME-001 - RPT-1001\tDRAFT"
    assert body.header.paragraphs[0].text == "ACME-001 - RPT-1001\tDRAFT"

    # No title-page bookmark found: front matter as a whole gets no page number.
    assert "PAGE" not in front_matter.footer.paragraphs[0]._p.xml
    assert "PAGE" in body.footer.paragraphs[0]._p.xml

    body_pg_num_type = body._sectPr.find(qn("w:pgNumType"))
    assert body_pg_num_type is not None
    assert body_pg_num_type.get(qn("w:start")) == "1"


def test_apply_layout_without_bookmarks_is_single_section(tmp_path):
    docx_path = tmp_path / "no_bookmark.docx"
    document = docx.Document()
    document.add_paragraph("Just one section.")
    document.save(str(docx_path))

    apply_layout(docx_path, header_left_text="Header Text", status="final")

    result = docx.Document(str(docx_path))
    assert len(result.sections) == 1
    assert result.sections[0].header.paragraphs[0].text == "Header Text\tFINAL"


def test_apply_layout_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        apply_layout("/does/not/exist.docx", header_left_text="x")


def _build_docx_with_crossref(path: Path) -> None:
    document = docx.Document()
    p = document.add_paragraph("See ")
    _add_ref_field_run(p, bookmark="Figure1", hyperlink=True)
    _add_page_field_run(document.sections[0].footer.paragraphs[0])
    document.save(str(path))


def _ref_instr_text(document) -> str:
    for instr in document.element.body.iter(qn("w:instrText")):
        if instr.text and "REF" in instr.text:
            return instr.text
    raise AssertionError("no REF field found")


def test_apply_layout_keeps_crossref_hyperlink_by_default(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    apply_layout(docx_path, header_left_text="Header")

    result = docx.Document(str(docx_path))
    assert "\\h" in _ref_instr_text(result)


def test_apply_layout_strips_crossref_hyperlink_when_disabled(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    apply_layout(docx_path, header_left_text="Header", crossref_hyperlinks=False)

    result = docx.Document(str(docx_path))
    text = _ref_instr_text(result)
    assert "\\h" not in text
    assert "REF Figure1" in text
    # PAGE fields are untouched -- only REF fields' \h switch is stripped.
    assert "PAGE" in result.sections[0].footer.paragraphs[0]._p.xml


def test_apply_layout_from_qmd_reads_crossref_hyperlinks_frontmatter(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    qmd_path = tmp_path / "report.qmd"
    qmd_path.write_text("---\ntitle: Foo\ncrossref-hyperlinks: false\n---\n\nBody.\n")

    apply_layout_from_qmd(docx_path, qmd_path, status="draft")

    result = docx.Document(str(docx_path))
    assert "\\h" not in _ref_instr_text(result)


def test_apply_layout_rejects_unrecognized_crossref_hyperlinks_value(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    with pytest.raises(LayoutError, match="crossref-hyperlinks"):
        apply_layout(docx_path, crossref_hyperlinks="every-other-tuesday")


def test_apply_layout_same_page_marks_ref_field_but_leaves_it_hyperlinked(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    apply_layout(docx_path, crossref_hyperlinks="same-page")

    result = docx.Document(str(docx_path))
    # The field itself is untouched -- still hyperlinked, the safe
    # fallback if same_page_crossrefs.resolve_same_page_crossrefs() never
    # runs against the filled document.
    assert "\\h" in _ref_instr_text(result)

    # A marker bookmark now sits immediately before the field, for that
    # later step to find.
    bookmark_names = [
        b.get(qn("w:name"))
        for b in result.element.body.iter(qn("w:bookmarkStart"))
        if b.get(qn("w:name"))
    ]
    assert any(name.startswith("quartifyr-crossref-target-") for name in bookmark_names)


def test_apply_layout_same_page_leaves_non_ref_fields_untouched(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    apply_layout(docx_path, crossref_hyperlinks="same-page")

    result = docx.Document(str(docx_path))
    assert "PAGE" in result.sections[0].footer.paragraphs[0]._p.xml
    assert result.sections[0].footer.paragraphs[0]._p.xml.count("PAGE") >= 1


def test_apply_layout_from_qmd_reads_same_page_crossref_hyperlinks(tmp_path):
    docx_path = tmp_path / "test.docx"
    _build_docx_with_crossref(docx_path)

    qmd_path = tmp_path / "report.qmd"
    qmd_path.write_text('---\ntitle: Foo\ncrossref-hyperlinks: "same-page"\n---\n\nBody.\n')

    apply_layout_from_qmd(docx_path, qmd_path, status="draft")

    result = docx.Document(str(docx_path))
    bookmark_names = [
        b.get(qn("w:name"))
        for b in result.element.body.iter(qn("w:bookmarkStart"))
        if b.get(qn("w:name"))
    ]
    assert any(name.startswith("quartifyr-crossref-target-") for name in bookmark_names)
