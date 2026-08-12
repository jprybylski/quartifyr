"""Generate a docx reference-template from a quartifyr StyleConfig.

The output is meant to be used as Quarto's ``reference-doc:`` when rendering
the report shell. Quarto/pandoc's docx writer and the ``quarto-plus``
extension key off specific built-in Word style *names* (``Title``,
``Heading 1``-``9``, ``Caption``, ...), so rather than inventing new style
names we mutate python-docx's bundled built-ins in place -- that keeps the
generated file a normal, valid reference-doc with no custom wiring required
on the Quarto side.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import BaseStyle

from .schema import StyleConfig

_PAGE_SIZES_IN = {
    "letter": (8.5, 11.0),
    "a4": (8.27, 11.69),
}


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#").upper())


def _set_font(style: BaseStyle, name: str, size_pt: float, *, color: str | None = None,
              bold: bool | None = None, italic: bool | None = None) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    if color is not None:
        font.color.rgb = _rgb(color)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic

    # East-Asian/complex-script run properties fall back to Word's own theme
    # fonts unless set explicitly, which produces a mismatched font in Word's
    # UI even though ASCII text looks right. Force them to match.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    # python-docx's bundled default template gives Title/Subtitle/Heading N
    # theme-relative font references (asciiTheme="majorHAnsi" etc, resolving
    # to the theme's major font -- Calibri in the default template's theme1.xml)
    # ALONGSIDE literal ascii/hAnsi attributes. When both are present on the
    # same rFonts element, Word/LibreOffice prefer the *theme* reference and
    # silently ignore the literal font -- confirmed: headings rendered in
    # Calibri (falling back to Arial where Calibri isn't installed) despite
    # font.name reading back as "Times New Roman" via python-docx's own API,
    # since that just reads the (ignored) literal attribute. Strip the theme
    # attributes so the literal font actually wins.
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rfonts.get(qn(theme_attr)) is not None:
            del rfonts.attrib[qn(theme_attr)]


def _set_paragraph_format(
    style: BaseStyle,
    *,
    line_spacing: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
    keep_with_next: bool | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    pf = style.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next
    if alignment is not None:
        pf.alignment = alignment


def _add_bottom_rule(style: BaseStyle, color: str) -> None:
    """Add a thin bottom border to every paragraph using ``style``."""
    ppr = style.element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color.lstrip("#").upper())
    pbdr.append(bottom)
    ppr.append(pbdr)


def _get_or_add_style(doc: DocumentObject, name: str, style_type: WD_STYLE_TYPE, base: str | None = None) -> BaseStyle:
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, style_type)
        if base is not None:
            style.base_style = doc.styles[base]
        return style


def _configure_toc_style(doc: DocumentObject, level: int, config: StyleConfig, page_width_in: float, margins_in: float) -> None:
    """Create/configure the 'TOC n' style Word applies to generated ToC/LOF/LOT entries."""
    name = f"TOC {level}"
    style = _get_or_add_style(doc, name, WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    _set_font(style, config.fonts.body, config.fonts.sizes.toc, color=config.colors.text)
    indent_in = 0.25 * (level - 1)
    style.paragraph_format.left_indent = Inches(indent_in)
    style.paragraph_format.space_after = Pt(4)

    # Right-aligned dot-leader tab stop at the text-area width, so entries
    # read as "Heading text ..... 4" the way a native Word ToC does.
    usable_width_in = page_width_in - (2 * margins_in)
    tab_stops = style.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(usable_width_in), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def _style_table_grid(doc: DocumentObject, config: StyleConfig) -> None:
    style = doc.styles["Table Grid"]
    _set_font(style, config.fonts.body, config.fonts.sizes.body, color=config.colors.text)

    tbl_pr = style.element.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        style.element.append(tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), config.colors.table_border.lstrip("#").upper())
        borders.append(border)
    # "Table Grid" already ships with a tblBorders element (auto-color) in
    # python-docx's default template; CT_TblPrBase only permits one, so swap
    # it in place rather than appending a second (invalid) one.
    existing_borders = tbl_pr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        existing_borders.addnext(borders)
        tbl_pr.remove(existing_borders)
    else:
        tbl_pr.append(borders)

    if config.table.header_bold:
        # Conditional formatting applied only to the header ("firstRow") band.
        style_pr = OxmlElement("w:tblStylePr")
        style_pr.set(qn("w:type"), "firstRow")
        rpr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rpr.append(b)
        style_pr.append(rpr)
        tc_pr = OxmlElement("w:tcPr")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), config.colors.table_header_fill.lstrip("#").upper())
        tc_pr.append(shd)
        style_pr.append(tc_pr)
        style.element.append(style_pr)


def _style_header(doc: DocumentObject, config: StyleConfig) -> None:
    """Styles the "Header" paragraph style so dynamic header text (added by
    quartifyr_styling.layout's post-render step, since the actual per-project
    text -- project code, report number, draft/final status -- isn't known
    until render time) inherits correct formatting automatically instead of
    Word's own default header look.
    """
    _set_font(doc.styles["Header"], config.fonts.body, config.fonts.sizes.footnote, color=config.colors.text)
    _set_paragraph_format(doc.styles["Header"], alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _add_page_number_footer(doc: DocumentObject, config: StyleConfig) -> None:
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style = doc.styles["Footer"]
    _set_font(doc.styles["Footer"], config.fonts.body, config.fonts.sizes.footnote, color=config.colors.text)

    if config.footer.text:
        run = footer_para.add_run(f"{config.footer.text}  ")
        run.font.name = config.fonts.body
        run.font.size = Pt(config.fonts.sizes.footnote)

    if not config.footer.show_page_number:
        return

    def _field_run(instr: str) -> None:
        run = footer_para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f" {instr} "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_sep)
        run._r.append(fld_end)

    _field_run("PAGE")


# Elements CT_Settings's schema (ECMA-376) requires to come *after*
# w:updateFields, in order -- copied from python-docx's own
# docx.oxml.settings.CT_Settings._tag_seq (not importable: that module
# deletes the name off the class after using it internally), starting
# right after "w:updateFields" itself. Used to find the correct insertion
# point among whatever settings python-docx's bundled default template
# (which this reference-doc is built on top of) happens to already
# contain -- appending blindly to the end would put it after elements the
# schema requires to precede it.
_SETTINGS_AFTER_UPDATE_FIELDS = (
    "w:hdrShapeDefaults",
    "w:footnotePr",
    "w:endnotePr",
    "w:compat",
    "w:docVars",
    "w:rsids",
    "m:mathPr",
    "w:attachedSchema",
    "w:themeFontLang",
    "w:clrSchemeMapping",
    "w:doNotIncludeSubdocsInStats",
    "w:doNotAutoCompressPictures",
    "w:forceUpgrade",
    "w:captions",
    "w:readModeInkLockDown",
    "w:smartTagType",
    "sl:schemaLibrary",
    "w:shapeDefaults",
    "w:doNotEmbedSmartTags",
    "w:decimalSymbol",
    "w:listSeparator",
)


def _enable_update_fields_on_open(doc: DocumentObject) -> None:
    """Sets ``<w:updateFields w:val="true"/>`` in the reference-doc's
    ``word/settings.xml``, so Word automatically recalculates every field
    (``TOC``, ``SEQ``, ``REF``, ``PAGE``, ...) the moment a delivered
    document is opened -- no manual "select all, F9", and no dependency on
    ``quartifyr-styling recalculate-fields``'s headless-LibreOffice
    automation, for anyone opening the file in real Microsoft Word.

    Not a full replacement for ``recalculate-fields``: LibreOffice doesn't
    reliably honor this flag the same way in every mode (confirmed only
    for interactive opens, not verified for the headless
    ``--convert-to``-style invocations this project's own tooling uses),
    and it can't help ``crossref-hyperlinks: "same-page"`` at all -- that
    mode's whole point is to bake in a one-time editorial decision before
    delivery, not leave it for whichever application opens the file to
    re-decide live. It's a complementary, zero-dependency win for the
    ordinary case of a human opening the final docx in Word, confirmed
    (via ``examples/demo-report/smoke_test.py``) to survive Quarto's
    reference-doc pipeline into the actual rendered output -- pandoc's
    docx writer carries the reference-doc's ``settings.xml`` through
    largely as-is.
    """
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")

    successor = next(
        (child for child in settings if child.tag in {qn(tag) for tag in _SETTINGS_AFTER_UPDATE_FIELDS}),
        None,
    )
    if successor is not None:
        successor.addprevious(update_fields)
    else:
        settings.append(update_fields)


def build_reference_docx(config: StyleConfig, output_path: str | Path) -> Path:
    """Build a Quarto ``reference-doc`` docx from ``config`` and write it to ``output_path``."""
    doc = Document()

    page_width_in, page_height_in = _PAGE_SIZES_IN[config.page.size]
    section = doc.sections[0]
    section.page_width = Inches(page_width_in)
    section.page_height = Inches(page_height_in)
    section.top_margin = Inches(config.page.margins_in.top)
    section.bottom_margin = Inches(config.page.margins_in.bottom)
    section.left_margin = Inches(config.page.margins_in.left)
    section.right_margin = Inches(config.page.margins_in.right)

    _set_font(doc.styles["Normal"], config.fonts.body, config.fonts.sizes.body, color=config.colors.text)
    _set_paragraph_format(
        doc.styles["Normal"],
        line_spacing=config.paragraph.line_spacing,
        space_after_pt=config.paragraph.space_after_pt,
    )

    _set_font(doc.styles["Title"], config.fonts.heading, config.fonts.sizes.title, color=config.colors.title, bold=True)
    _set_paragraph_format(doc.styles["Title"], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    if config.title_page.show_rule_under_title:
        _add_bottom_rule(doc.styles["Title"], config.colors.rule)

    _set_font(doc.styles["Subtitle"], config.fonts.heading, config.fonts.sizes.subtitle, color=config.colors.title, italic=True)
    _set_paragraph_format(doc.styles["Subtitle"], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)

    for level in range(1, 7):
        style = doc.styles[f"Heading {level}"]
        _set_font(
            style,
            config.fonts.heading,
            config.fonts.sizes.heading.get(level),
            color=config.colors.heading,
            bold=config.heading.bold,
        )
        _set_paragraph_format(
            style,
            space_before_pt=config.heading.space_before_pt,
            space_after_pt=config.heading.space_after_pt,
            keep_with_next=config.heading.keep_with_next,
        )

    _set_font(doc.styles["Caption"], config.fonts.body, config.fonts.sizes.caption, color=config.colors.caption, bold=False, italic=True)
    _set_paragraph_format(doc.styles["Caption"], space_before_pt=4, space_after_pt=8)

    for level in (1, 2, 3):
        _configure_toc_style(doc, level, config, page_width_in, config.page.margins_in.left)

    _style_table_grid(doc, config)
    _style_header(doc, config)
    _add_page_number_footer(doc, config)
    _enable_update_fields_on_open(doc)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
