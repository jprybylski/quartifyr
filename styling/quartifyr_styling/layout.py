"""Post-render layout: page-numbering restart and a dynamic header/footer.

Splits the rendered docx into OOXML sections at two bookmarks:
``quartifyr-front-matter-start`` (emitted automatically by
``title_page.lua`` right after the title page) and
``quartifyr-body-start`` (``{{< body-start >}}`` in the shell ``.qmd``,
see ``_extensions/quartifyr/body_start.lua``). With both present, that's
three sections: the title page alone, the rest of the front matter (ToC,
list of figures/tables, abbreviations, synopsis, signature pages, ...),
and the body. With only ``quartifyr-body-start`` present (no title page
was rendered), that's two: front matter and body.

Header (opt-in via ``header-format:``): two zones -- a resolved
``header-format:`` template on the left, and the draft/final status on
the right (always shown once a header is enabled, mirroring
``title_page.lua``'s own always-shown status stamp). Applies to every
page, including the title page.

Footer: two zones -- an optional confidentiality label on the left
(reusing the ``confidentiality:`` field ``title_page.lua`` already renders
on the title page), and a page number on the right. The whole front
matter numbers in lowercase roman, starting at "i" on the title page
itself (the title page keeps its own OOXML section -- see below -- but
that's no longer about hiding its page number; it's now section 1 of the
same roman sequence the rest of the front matter continues as section 2,
starting at "ii"). The body section restarts at arabic "1".

Why three real OOXML sections rather than one front-matter section with
OOXML's "different first page" mechanism (``w:titlePg``): that was the
first approach tried here, and it round-trips visibly broken through real
Microsoft Word -- opening, recalculating fields, and saving collapsed the
distinct first-page footer part into the regular one (reproduced
directly: a 3-header/3-footer-part file came back with only 2 of each,
and the first-page-specific content was gone). The plain "clone a section
break's sectPr at a bookmark, strip its header/footerReference" technique
below is the same one already proven to round-trip correctly through Word
for the front-matter/body split, so the title page just gets one more of
the same kind of break rather than a different, apparently less robust,
mechanism.

Why this is a Python post-processing step and not more Lua: a second,
independent header/footer means adding new *parts* to the docx package
(``word/header2.xml``, relationships, content-type overrides) -- package-
level plumbing a Lua filter's ``RawBlock`` injection can't do (it can only
inject XML into the existing body flow). ``python-docx``'s ``Section`` API
handles that plumbing correctly (confirmed via testing), so this runs
after Quarto's render as a distinct step, via ``quartifyr-styling
apply-layout``.

Also controls crossref hyperlinking via ``crossref-hyperlinks:``: every
figure/table/appendix cross-reference -- whether from quarto-plus's
``crossref`` shortcode or this extension's own ``appendix_crossref`` --
compiles to a Word ``REF <bookmark> \\h`` field, and the ``\\h`` switch is
what makes it a clickable hyperlink. Three modes, since the switch is
identical regardless of which extension emitted the field and is rewritten
once here at the OOXML level rather than in each Lua shortcode separately:

- ``true`` (default): leave every crossref hyperlinked, matching
  quarto-plus's own out-of-the-box behavior.
- ``false``: strip ``\\h`` from every crossref, document-wide -- never a
  hyperlink, regardless of where the reference and its target land.
- ``"same-page"``: the org convention this exists for -- a crossref is a
  hyperlink only when its target is on a *different* page; on the same
  page it renders as plain resolved text. This can't be resolved here:
  real page numbers don't exist yet at this step (pass 1 -- the shell is
  still empty of reportifyr's actual content, so pagination here would be
  meaningless even if computed). An earlier version of this mode tried
  pushing the comparison into a live nested Word field
  (``IF {PAGE} = {PAGEREF bookmark} "{REF bookmark}" "{REF bookmark
  \\h}"``) so Word/LibreOffice would resolve it whenever the *final*
  document gets paginated -- abandoned after confirming, via a real
  headless-LibreOffice round-trip, that it silently mangles the nested
  field into garbage (a stray always-hyperlinked ``PAGEREF`` plus a
  broken ``IF FORMULA "" ""``) rather than evaluating it -- not just
  unverified, but demonstrated broken through the one engine available to
  test with. Instead, this mode only *marks* each same-page-mode crossref
  here (a small bookmark next to its ``REF`` field, so it can be found
  again later) and leaves the field itself hyperlinked, matching mode
  ``true``, as the safe fallback if nothing ever resolves the mark. The
  actual page-number comparison and ``\\h`` decision happens afterward, in
  ``quartifyr_styling.same_page_crossrefs.resolve_same_page_crossrefs()``
  -- which must run after reportifyr's pass 2 (real content, real
  pagination) -- see that module for why and how.
"""

from __future__ import annotations

import copy
import re
from pathlib import Path

import docx
import yaml
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu

from ._ooxml_fields import (
    SAME_PAGE_MARKER_ID_BASE,
    SAME_PAGE_MARKER_PREFIX,
    match_ref_field_run_group,
    strip_hyperlink_switch,
)

_FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?\n)---\s*\n", re.DOTALL)
_FRONT_MATTER_START_BOOKMARK = "quartifyr-front-matter-start"
_BODY_START_BOOKMARK = "quartifyr-body-start"


class LayoutError(RuntimeError):
    """Raised when applying the post-render layout fails."""


def read_qmd_frontmatter(qmd_path: str | Path) -> dict:
    """Parses the YAML frontmatter block (between the first two ``---`` lines) of a ``.qmd`` file."""
    text = Path(qmd_path).read_text(encoding="utf-8")
    match = _FRONTMATTER_RE.match(text)
    if not match:
        return {}
    return yaml.safe_load(match.group(1)) or {}


def resolve_header_left_text(frontmatter: dict, status: str) -> str | None:
    """Resolves ``header-format:``'s ``{placeholder}`` template against frontmatter values plus ``status``.

    Placeholders reference frontmatter keys directly (e.g. ``{project}``
    reads ``frontmatter["project"]``); a ``{status}`` placeholder is also
    available (``DRAFT``/``FINAL``, from the ``status`` argument -- not
    read from frontmatter, since the resolved render status can differ
    from whatever static ``document-status:`` the ``.qmd`` happens to
    have) for templates that want it inline, though the header's right
    zone already always shows status on its own -- see ``apply_layout()``.

    Returns ``None`` if no ``header-format:`` is set (the header feature
    is opt-in).
    """
    template = frontmatter.get("header-format")
    if not template:
        return None
    template = str(template)

    context = {"status": status.upper()}
    for key, value in frontmatter.items():
        if isinstance(value, (str, int, float)):
            context[key] = value

    try:
        return template.format(**context)
    except (KeyError, IndexError) as exc:
        raise LayoutError(
            f"header-format template {template!r} references a placeholder not found in frontmatter: {exc}"
        ) from exc


def resolve_confidential_label(frontmatter: dict) -> str:
    """Resolves the footer's left-side confidentiality label.

    Reuses ``confidentiality:`` -- the same field ``title_page.lua`` already
    renders on the title page (e.g. ``"Confidential — Do Not Distribute"``)
    -- rather than a second, separately-configured field: whatever that
    field is set to appears verbatim on every footer too, and it's omitted
    entirely (``""``) when ``confidentiality:`` isn't set.
    """
    value = frontmatter.get("confidentiality")
    if not value:
        return ""
    return str(value)


def _add_page_field(paragraph) -> None:
    """Appends a live PAGE field to ``paragraph``, across five runs matching
    what Word itself produces when a field is inserted (``begin``,
    ``instrText``, ``separate``, a cached placeholder result, ``end``).

    The placeholder result text (a literal "1") between ``separate`` and
    ``end`` is load-bearing, not cosmetic: leaving it out (a single run
    with all four field-character elements and no result content, tried
    first) rendered visibly wrong in real Word before any field
    recalculation happened -- specifically, the field failed to honor the
    paragraph's right tab stop, landing well short of the right margin
    instead. A field with no measurable content at layout time apparently
    can't be correctly positioned against a tab stop; giving it a real,
    if provisional, glyph fixes that, and Word replaces it with the
    correct page number on the next recalculation (or immediately, since
    Word auto-calculates simple fields like PAGE on open).
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run._r.append(instr_text)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    paragraph.add_run("1")

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)


def _content_width(section) -> Emu:
    return Emu(section.page_width - section.left_margin - section.right_margin)


def _clear_inherited_tab_stops(paragraph, *, keep_positions=()) -> None:
    """Cancels every tab stop the paragraph's style defines, except any in
    ``keep_positions`` (about to be explicitly redefined anyway, so
    clearing and immediately re-adding at that exact position would just
    emit a redundant, ambiguous same-position pair of ``<w:tab>``
    elements).

    Load-bearing, confirmed via a real Word render: python-docx's default
    "Header"/"Footer" styles (which ``build_template.py``'s reference-doc
    inherits) already define their own center/right tabs (Word's stock
    3.25in/6.5in positions, for a standard 1in-margin letter page) --
    OOXML *merges* a paragraph's own ``w:tabs`` with its style's rather
    than replacing them. With both a right tab at the true content width
    and the style's center tab still active, a `<w:tab/>` character
    resolved to whichever tab stop came first after the cursor -- the
    style's inherited center tab, not the paragraph's own right tab --
    so a page number ended up landing near mid-page instead of flush
    right. Explicitly clearing (``w:val="clear"``) each inherited
    position before adding new ones removes the stale ones instead of
    layering on top of them.
    """
    style = paragraph.style
    if style is None:
        return
    inherited_positions = [ts.position for ts in style.paragraph_format.tab_stops]
    tab_stops = paragraph.paragraph_format.tab_stops
    for position in inherited_positions:
        if position in keep_positions:
            continue
        tab_stops.add_tab_stop(position, WD_TAB_ALIGNMENT.CLEAR)


def _write_header_paragraph(paragraph, section, left_text: str, right_text: str) -> None:
    """``left_text`` left-aligned, ``right_text`` right-aligned."""
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    width = _content_width(section)
    _clear_inherited_tab_stops(paragraph, keep_positions={width})
    paragraph.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)

    paragraph.add_run(left_text)
    run = paragraph.add_run("")
    run.add_tab()
    paragraph.add_run(right_text)


def _write_footer_paragraph(paragraph, section, left_text: str, add_page_field: bool) -> None:
    """``left_text`` left-aligned, an optional live PAGE field right-aligned."""
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    width = _content_width(section)
    _clear_inherited_tab_stops(paragraph, keep_positions={width})
    paragraph.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)

    paragraph.add_run(left_text)
    if add_page_field:
        run = paragraph.add_run("")
        run.add_tab()
        _add_page_field(paragraph)


def _set_pg_num_type(section, *, start: int, fmt: str | None = None) -> None:
    sectPr = section._sectPr
    existing = sectPr.find(qn("w:pgNumType"))
    if existing is not None:
        sectPr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    if fmt:
        pg_num_type.set(qn("w:fmt"), fmt)
    sectPr.append(pg_num_type)


_VALID_CROSSREF_HYPERLINK_MODES = ("always", "never", "same-page")


def _resolve_crossref_hyperlink_mode(value: bool | str) -> str:
    """Normalizes the ``crossref-hyperlinks:`` frontmatter value (a YAML
    ``true``/``false``, or the literal string ``"same-page"``) into one of
    ``_VALID_CROSSREF_HYPERLINK_MODES``.
    """
    if isinstance(value, bool):
        return "never" if value is False else "always"
    if isinstance(value, str) and value.strip().lower() == "same-page":
        return "same-page"
    raise LayoutError(f'crossref-hyperlinks: {value!r} not recognized -- use `true`, `false`, or `"same-page"`')


def _strip_crossref_hyperlinks(document) -> None:
    """Removes the ``\\h`` switch from every ``REF`` field's instruction
    text, so Word renders the resolved cross-reference as plain text
    instead of a hyperlink.

    Field-agnostic on purpose: it walks raw ``w:instrText`` elements
    rather than calling into any particular Lua shortcode's output, so it
    catches quarto-plus's ``crossref`` (figures/tables) and this
    extension's own ``appendix_crossref`` identically.
    """
    for instr_text in document.element.body.iter(qn("w:instrText")):
        strip_hyperlink_switch(instr_text)


def _mark_crossrefs_for_same_page_resolution(document) -> None:
    """Inserts a small marker bookmark immediately before every ``REF
    <bookmark> [\\h]`` field's run group, so
    ``same_page_crossrefs.resolve_same_page_crossrefs()`` can find each one
    again after reportifyr's pass 2 (once real pagination exists) and
    decide whether it should be hyperlinked.

    Doesn't touch the field itself -- it's left hyperlinked, the same as
    mode ``true``, as the safe fallback if that later resolution step
    never runs.
    """
    marker_id = SAME_PAGE_MARKER_ID_BASE
    marker_count = 0
    for p in document.element.body.iter(qn("w:p")):
        runs = list(p.findall(qn("w:r")))
        i = 0
        while i < len(runs):
            found = match_ref_field_run_group(runs, i)
            if found is None:
                i += 1
                continue
            marker_count += 1
            marker_id += 1
            anchor = runs[i]
            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), str(marker_id))
            bookmark_start.set(qn("w:name"), f"{SAME_PAGE_MARKER_PREFIX}{marker_count}")
            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), str(marker_id))
            anchor.addprevious(bookmark_start)
            anchor.addprevious(bookmark_end)
            i += 1


def _find_bookmark_paragraph(document, name: str):
    body = document.element.body
    for p in body.iter(qn("w:p")):
        for bookmark in p.iter(qn("w:bookmarkStart")):
            if bookmark.get(qn("w:name")) == name:
                return p
    return None


def _insert_section_break(bookmark_p, final_sectPr) -> None:
    """Ends a new OOXML section at ``bookmark_p`` by cloning ``final_sectPr``
    into its ``pPr``, with any header/footerReference stripped from the
    clone first.

    Stripping the references is load-bearing, confirmed via python-docx's
    own source (``docx/section.py``): ``is_linked_to_previous`` treats
    "has a header/footerReference element" as *already independently
    defined*, and its setter is a no-op if that's already true --
    "assigning False ... only [creates a new definition] if no definition
    is already present". The reference-doc's footer (``build_template.py``
    's page-number footer) means an unstripped clone starts out WITH a
    footerReference, so sections would silently end up permanently sharing
    the exact same footer part no matter how ``is_linked_to_previous`` is
    called afterward -- reproduced repeatedly before finding this.
    Stripping the reference makes the new section start in the actual
    "inherited" state python-docx expects, so a later unlink-then-
    customize actually creates an independent part.
    """
    new_sectPr = copy.deepcopy(final_sectPr)
    for ref_tag in ("w:headerReference", "w:footerReference"):
        for ref in new_sectPr.findall(qn(ref_tag)):
            new_sectPr.remove(ref)
    pPr = bookmark_p.get_or_add_pPr()
    pPr.append(new_sectPr)


def apply_layout(
    docx_path: str | Path,
    *,
    header_left_text: str | None = None,
    status: str = "DRAFT",
    confidential_label: str = "",
    show_page_numbers: bool = True,
    crossref_hyperlinks: bool | str = True,
) -> Path:
    """Applies header/footer + page-numbering layout to ``docx_path``, in place.

    If no ``quartifyr-body-start`` bookmark is found (the shell ``.qmd``
    never used ``{{< body-start >}}``), the document is treated as a
    single section: the header (if ``header_left_text`` is given)
    applies throughout, and no footer/page-numbering changes happen
    (there's no defined split point for either).

    ``crossref_hyperlinks`` (default ``True``, matching quarto-plus's own
    always-hyperlinked ``REF ... \\h`` fields): ``True``/``False`` for
    always/never hyperlinked, document-wide, or the string ``"same-page"``
    to mark every crossref for later resolution by
    ``same_page_crossrefs.resolve_same_page_crossrefs()`` once real
    pagination exists (post reportifyr) -- see the module docstring for
    how each mode works.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"docx not found: {docx_path}")
    crossref_hyperlink_mode = _resolve_crossref_hyperlink_mode(crossref_hyperlinks)

    document = docx.Document(str(docx_path))
    front_matter_start_p = _find_bookmark_paragraph(document, _FRONT_MATTER_START_BOOKMARK)
    body_start_p = _find_bookmark_paragraph(document, _BODY_START_BOOKMARK)

    if body_start_p is not None:
        final_sectPr = document.element.body.find(qn("w:sectPr"))
        if final_sectPr is None:
            raise LayoutError(f"{docx_path}: no trailing sectPr found in document body")

        has_title_section = front_matter_start_p is not None
        if has_title_section:
            _insert_section_break(front_matter_start_p, final_sectPr)
        _insert_section_break(body_start_p, final_sectPr)

        # Re-open so python-docx's `sections` (computed by scanning the
        # XML) picks up the newly-inserted section(s); the in-memory
        # object graph built at Document() time doesn't refresh on its own
        # after a raw XML insert like the one(s) just above.
        document.save(str(docx_path))
        document = docx.Document(str(docx_path))

        if has_title_section:
            title_section, front_matter_section, body_section = document.sections[:3]
            all_sections = [title_section, front_matter_section, body_section]
        else:
            title_section = None
            front_matter_section, body_section = document.sections[:2]
            all_sections = [front_matter_section, body_section]

        # IMPORTANT: unlink every section's header/footer *before* setting
        # any content on any of them. Interleaving unlink-then-modify per
        # section (unlink one, modify it, unlink the next, modify it) was
        # confirmed via a real end-to-end test to make sections end up
        # sharing a single part -- whatever a later section's unlink did
        # ended up mutating an earlier section's part in place rather than
        # creating its own. Doing all the unlinking first, matching the
        # exact sequence validated in isolated testing, avoids that.
        if header_left_text is not None:
            for section in all_sections:
                section.header.is_linked_to_previous = False
        for section in all_sections:
            section.footer.is_linked_to_previous = False

        if header_left_text is not None:
            for section in all_sections:
                _write_header_paragraph(section.header.paragraphs[0], section, header_left_text, status.upper())

        if title_section is not None:
            # Title page: lowercase roman, starting at "i".
            _write_footer_paragraph(
                title_section.footer.paragraphs[0],
                title_section,
                confidential_label,
                add_page_field=show_page_numbers,
            )
            if show_page_numbers:
                _set_pg_num_type(title_section, start=1, fmt="lowerRoman")
            # Rest of front matter: same roman sequence, continuing at "ii".
            _write_footer_paragraph(
                front_matter_section.footer.paragraphs[0],
                front_matter_section,
                confidential_label,
                add_page_field=show_page_numbers,
            )
            if show_page_numbers:
                _set_pg_num_type(front_matter_section, start=2, fmt="lowerRoman")
        else:
            # No title-page bookmark (title_page.lua didn't run): front
            # matter as a whole gets no page number, matching the
            # single-front-matter-section behavior before the title page
            # was split out on its own.
            _write_footer_paragraph(
                front_matter_section.footer.paragraphs[0], front_matter_section, confidential_label, add_page_field=False
            )

        # Body: a fresh, explicit PAGE-field footer, restarting at 1 --
        # NOT relying on "inherit whatever footer the reference-doc set
        # up", since a working footer part can't be assumed to survive
        # the unlink dance above; building it directly here removes that
        # dependency entirely.
        _write_footer_paragraph(
            body_section.footer.paragraphs[0], body_section, confidential_label, add_page_field=show_page_numbers
        )
        if show_page_numbers:
            _set_pg_num_type(body_section, start=1, fmt="decimal")
    elif header_left_text is not None:
        # No body-start marker: single section, header applies throughout,
        # no footer/page-numbering changes (there's no defined split point
        # for either).
        section = document.sections[0]
        section.header.is_linked_to_previous = False
        _write_header_paragraph(section.header.paragraphs[0], section, header_left_text, status.upper())

    if crossref_hyperlink_mode == "never":
        _strip_crossref_hyperlinks(document)
    elif crossref_hyperlink_mode == "same-page":
        _mark_crossrefs_for_same_page_resolution(document)

    document.save(str(docx_path))
    return docx_path


def apply_layout_from_qmd(
    docx_path: str | Path,
    qmd_path: str | Path,
    *,
    status: str,
) -> Path:
    """Convenience wrapper: reads ``qmd_path``'s frontmatter, resolves
    ``header-format:`` against it and ``status``, resolves the footer's
    confidentiality label from ``confidentiality:``, resolves
    ``crossref-hyperlinks:`` (default ``True`` -- ``True``/``False``/
    ``"same-page"``), and applies the layout.
    """
    frontmatter = read_qmd_frontmatter(qmd_path)
    header_left_text = resolve_header_left_text(frontmatter, status)
    confidential_label = resolve_confidential_label(frontmatter)
    crossref_hyperlinks = frontmatter.get("crossref-hyperlinks", True)
    return apply_layout(
        docx_path,
        header_left_text=header_left_text,
        status=status,
        confidential_label=confidential_label,
        crossref_hyperlinks=crossref_hyperlinks,
    )
